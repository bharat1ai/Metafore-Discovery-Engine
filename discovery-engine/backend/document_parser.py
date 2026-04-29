import io
from pathlib import Path


def parse_document(filename: str, content: bytes) -> str:
    """Parse PDF, DOCX, or TXT and return plain text."""
    ext = Path(filename).suffix.lower()

    if ext == ".txt":
        return content.decode("utf-8", errors="ignore")

    if ext == ".pdf":
        import fitz  # PyMuPDF

        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text() for page in doc]
        return "\n\n".join(pages)

    if ext in (".docx", ".doc"):
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # python-docx's doc.paragraphs iterator skips text inside tables. Without this
        # block, table-heavy enterprise documents (registers, role matrices, SLA tables)
        # would be invisible to extraction. We append cell text after the body paragraphs
        # — graph extraction is order-agnostic.
        for table in doc.tables:
            for row in table.rows:
                row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")
