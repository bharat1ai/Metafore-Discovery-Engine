"""Generate professional DOCX versions of the two banking demo documents."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent
BRAND      = RGBColor(0x03, 0x68, 0x68)   # #036868
BRAND_DARK = RGBColor(0x02, 0x4F, 0x4F)   # #024F4F
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
RED        = RGBColor(0xDC, 0x26, 0x26)
GREEN      = RGBColor(0x16, 0xA3, 0x4A)
GREY_LIGHT = RGBColor(0xF3, 0xF4, 0xF6)
GREY_MID   = RGBColor(0x6B, 0x72, 0x80)


# ── Helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb):
    # rgb is RGBColor (str(rgb) returns 6-char hex) or a plain hex string
    hex_color = str(rgb) if not isinstance(rgb, str) else rgb
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def cell_para(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.paragraphs[0].alignment = align
    return cell.paragraphs[0]


def add_header_banner(doc, title_line1, title_line2, ref, version, date, owner):
    """Dark teal header banner with document identity."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, BRAND_DARK)
    cell.width = Inches(6.5)
    p1 = cell.paragraphs[0]
    p1.clear()
    r1 = p1.add_run(title_line1)
    r1.font.color.rgb = WHITE
    r1.font.size = Pt(9)
    r1.font.bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = cell.add_paragraph()
    r2 = p2.add_run(title_line2)
    r2.font.color.rgb = WHITE
    r2.font.size = Pt(16)
    r2.font.bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = cell.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"{ref}   |   {version}   |   Effective {date}")
    mr.font.color.rgb = RGBColor(0xA7, 0xD4, 0xD4)
    mr.font.size = Pt(8)

    p3 = cell.add_paragraph()
    r3 = p3.add_run(f"Owner: {owner}   |   Classification: Internal Use Only")
    r3.font.color.rgb = RGBColor(0xA7, 0xD4, 0xD4)
    r3.font.size = Pt(8)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_DARK
    # underline via border would need XML; use simple bold colour instead
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND
    return p


def body(doc, text, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x2D)
    return p


def meta_table(doc, rows_data):
    """Two-column label/value table for stage metadata."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    col_w = [Inches(1.6), Inches(4.9)]
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].width = col_w[0]
        row.cells[1].width = col_w[1]
        set_cell_bg(row.cells[0], RGBColor(0xE6, 0xF4, 0xF4))
        cell_para(row.cells[0], label, bold=True, size=9, color=BRAND_DARK)
        cell_para(row.cells[1], value, size=9)
    doc.add_paragraph()
    return table


def status_badge(para, status_text, color):
    run = para.add_run(f"  {status_text}  ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = color


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — SOP
# ══════════════════════════════════════════════════════════════════════

def build_sop():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    add_header_banner(
        doc,
        "COMMERCIAL BANKING DIVISION",
        "Loan Origination Standard Operating Procedure",
        "SOP-LO-004.2",
        "Version 4.2",
        "1 March 2024",
        "Head of Commercial Banking",
    )

    # ── 1. Purpose and Scope
    heading1(doc, "1.  Purpose and Scope")
    body(doc,
        "This Standard Operating Procedure governs the end-to-end loan origination process "
        "for all commercial loan applications processed by the Commercial Banking Division. "
        "It applies to all branch locations and covers applications from initial enquiry "
        "through to commitment letter issuance."
    )
    body(doc,
        "All Loan Processing Officers, Credit Analysts, Branch Managers, and Credit Committee "
        "members are required to comply with this SOP. Non-compliance must be escalated to the "
        "Compliance Officer within 24 hours of identification."
    )

    # ── 2. Process Overview
    heading1(doc, "2.  Process Overview")
    body(doc,
        "The loan origination process consists of eight sequential stages. Each stage has a "
        "designated responsible role, a system of record, and a mandatory completion gate "
        "before the next stage may begin."
    )

    stages = [
        ("Stage 1", "AML and Compliance Pre-Screening"),
        ("Stage 2", "Credit Bureau Pull"),
        ("Stage 3", "Document Collection and Upload"),
        ("Stage 4", "Debt-to-Income Ratio Calculation"),
        ("Stage 5", "Property Appraisal"),
        ("Stage 6", "Credit Assessment and Underwriting"),
        ("Stage 7", "Sanctioning Decision"),
        ("Stage 8", "Commitment Letter Issuance"),
    ]
    table = doc.add_table(rows=1 + len(stages), cols=2)
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    set_cell_bg(hdr.cells[0], BRAND)
    set_cell_bg(hdr.cells[1], BRAND)
    cell_para(hdr.cells[0], "Stage",       bold=True, color=WHITE, size=9)
    cell_para(hdr.cells[1], "Description", bold=True, color=WHITE, size=9)
    for i, (s, d) in enumerate(stages):
        r = table.rows[i + 1]
        bg = RGBColor(0xF0, 0xFA, 0xFA) if i % 2 == 0 else WHITE
        set_cell_bg(r.cells[0], bg)
        set_cell_bg(r.cells[1], bg)
        cell_para(r.cells[0], s, bold=True, size=9, color=BRAND_DARK)
        cell_para(r.cells[1], d, size=9)
    doc.add_paragraph()

    # ── 3. Detailed Procedures
    heading1(doc, "3.  Detailed Procedures")

    stage_data = [
        {
            "title": "3.1  Stage 1 — AML and Compliance Pre-Screening",
            "meta": [
                ("Responsible Role", "Loan Processing Officer"),
                ("System of Record", "Compliance Tracker"),
                ("SLA",              "Completed within 4 business hours of application receipt"),
                ("Gate",             "AML Cleared status must be recorded in Compliance Tracker before any credit enquiry is initiated"),
            ],
            "body": [
                "The Loan Processing Officer must initiate an Anti-Money Laundering screen through the "
                "Compliance Tracker system immediately upon receipt of a new loan application. The AML "
                "check covers sanctions list screening, Politically Exposed Person (PEP) checks, and "
                "adverse media review.",
                "The AML pre-screen MUST be completed and show a Cleared status in the Compliance "
                "Tracker before the credit bureau pull in Stage 2 may be initiated. This sequence is "
                "mandatory under the Bank's AML Policy (POL-AML-001). Under no circumstances may a "
                "credit bureau query be submitted while an AML screen is pending or shows an Alert status.",
                "If the AML screen returns an Alert, the application must be referred to the Compliance "
                "Officer before proceeding.",
            ],
        },
        {
            "title": "3.2  Stage 2 — Credit Bureau Pull",
            "meta": [
                ("Responsible Role", "Loan Processing Officer"),
                ("System of Record", "LoanIQ"),
                ("SLA",              "Completed within 1 business day of AML clearance"),
                ("Gate",             "Bureau reports filed in Document Management System"),
            ],
            "body": [
                "The Loan Processing Officer initiates a dual credit bureau pull through LoanIQ using "
                "both Experian and TransUnion. Both bureau reports must be formally retrieved and filed "
                "in the designated folder structure within the Document Management System.",
                "Bureau reports must be obtained prior to the commencement of any credit analysis "
                "activity. The LoanIQ system records timestamps of all bureau queries and these are "
                "reviewed during compliance audits.",
            ],
        },
        {
            "title": "3.3  Stage 3 — Document Collection and Upload",
            "meta": [
                ("Responsible Role", "Loan Processing Officer"),
                ("System of Record", "Document Management System"),
                ("SLA",              "All required documents uploaded within 48 hours of application receipt (80% compliance threshold)"),
                ("Gate",             "Document Completeness Certificate issued and countersigned by Branch Manager"),
            ],
            "body": [
                "The Loan Processing Officer is responsible for collecting all required customer "
                "documentation and uploading it to the Document Management System within 48 hours "
                "of application receipt.",
                "Required documents include: signed loan application form, proof of identity, proof "
                "of address, financial statements (3 years), business plan (commercial loans), and "
                "property details where applicable.",
                "Upon completion of document upload, the Loan Processing Officer must issue a Document "
                "Completeness Certificate and obtain countersignature from the Branch Manager. This "
                "certificate is a mandatory prerequisite for Stage 5 (Property Appraisal). Under no "
                "circumstances may an appraisal be commissioned before the Document Completeness "
                "Certificate has been issued and countersigned.",
            ],
        },
        {
            "title": "3.4  Stage 4 — Debt-to-Income Ratio Calculation",
            "meta": [
                ("Responsible Role", "Credit Analyst"),
                ("System of Record", "Loan Application Worksheet"),
                ("SLA",              "Completed within 1 business day of document upload"),
                ("Gate",             "DTI calculation certified on Loan Application Worksheet"),
            ],
            "body": [
                "The Debt-to-Income ratio calculation must be performed and formally certified by a "
                "Credit Analyst on the standard Loan Application Worksheet. Only personnel holding "
                "the Credit Analyst role designation within LoanIQ may certify the DTI calculation "
                "field. Loan Processing Officers and Loan Officers do not have authority to certify "
                "DTI calculations under this SOP.",
                "The DTI ratio is calculated as total monthly debt obligations divided by gross monthly "
                "income. Applications with a DTI ratio exceeding 43% are referred to the Senior "
                "Underwriter for further assessment before proceeding to Credit Assessment.",
            ],
        },
        {
            "title": "3.5  Stage 5 — Property Appraisal",
            "meta": [
                ("Responsible Role", "Loan Processing Officer"),
                ("System of Record", "Appraisal Management System"),
                ("SLA",              "Appraisal report received within 5 business days of appraisal instruction"),
                ("Gate",             "Document Completeness Certificate (from Stage 3) must exist before appraisal vendor is engaged"),
            ],
            "body": [
                "The Loan Processing Officer engages an approved appraisal vendor through the Appraisal "
                "Management System. The system will only permit appraisal instruction if a valid Document "
                "Completeness Certificate exists for the application. A property appraisal may not be "
                "commissioned until Stage 3 is complete.",
                "The appraisal report must be filed in the Document Management System upon receipt and "
                "linked to the relevant application record in LoanIQ.",
            ],
        },
        {
            "title": "3.6  Stage 6 — Credit Assessment and Underwriting",
            "meta": [
                ("Responsible Role", "Senior Underwriter (DTI > 43%) / Credit Analyst (standard)"),
                ("System of Record", "LoanIQ"),
                ("SLA",              "Assessment completed within 2 business days of all Stage 3–5 inputs being available"),
                ("Gate",             "Credit Assessment Report filed in LoanIQ"),
            ],
            "body": [
                "The Credit Analyst or Senior Underwriter conducts a comprehensive credit assessment "
                "covering: DTI analysis, credit bureau report review, financial statement analysis, "
                "collateral valuation, and sector risk evaluation.",
                "The completed Credit Assessment Report must be filed in LoanIQ and forms the primary "
                "input for the Sanctioning Decision in Stage 7.",
            ],
        },
        {
            "title": "3.7  Stage 7 — Sanctioning Decision",
            "meta": [
                ("Responsible Role", "Credit Committee"),
                ("System of Record", "Credit Committee System"),
                ("SLA",              "Decision issued within 1 business day of Credit Assessment Report submission"),
                ("Escalation",       "Branch Manager must escalate any application exceeding 2 business days without a decision"),
            ],
            "body": [
                "The Credit Committee reviews the Credit Assessment Report and issues a sanctioning "
                "decision: Approved, Conditionally Approved, or Declined.",
                "The 1 business day SLA is mandatory. Branch Managers are required to monitor "
                "sanctioning queue status daily via the LoanIQ dashboard and initiate escalation to "
                "the Credit Committee for any application that has not received a decision within 2 "
                "business days.",
                "If the Credit Committee does not convene within the escalation window, the Branch "
                "Manager must notify the Head of Commercial Banking.",
            ],
        },
        {
            "title": "3.8  Stage 8 — Commitment Letter Issuance",
            "meta": [
                ("Responsible Role", "Branch Manager"),
                ("System of Record", "Document Management System"),
                ("SLA",              "Issued within 1 business day of sanctioning approval"),
                ("Gate",             "Sanctioning decision must be Approved or Conditionally Approved"),
            ],
            "body": [
                "Upon receipt of an Approved or Conditionally Approved sanctioning decision from the "
                "Credit Committee, the Branch Manager issues the formal Commitment Letter to the "
                "customer. The letter must include all approved terms, conditions, fee disclosures, "
                "and expiry date.",
                "The Branch Manager is the sole signatory for Commitment Letters. Letters must be "
                "filed in the Document Management System and linked to the application in LoanIQ.",
            ],
        },
    ]

    for sd in stage_data:
        heading2(doc, sd["title"])
        meta_table(doc, sd["meta"])
        for para in sd["body"]:
            body(doc, para)

    # ── 4. Roles and Responsibilities
    heading1(doc, "4.  Roles and Responsibilities")
    roles = [
        ("Loan Processing Officer", "Stages 1, 2, 3, 5 — initiates AML screening, pulls bureau reports, collects and uploads documents, engages appraisal vendors"),
        ("Credit Analyst",          "Stages 4, 6 — certifies DTI calculations, conducts standard credit assessments"),
        ("Senior Underwriter",      "Stage 6 — conducts credit assessment for high-DTI applications (DTI > 43%)"),
        ("Branch Manager",          "Stage 3 (countersign), Stage 7 (escalation monitoring), Stage 8 (commitment letter signatory)"),
        ("Credit Committee",        "Stage 7 — reviews Credit Assessment Reports and issues sanctioning decisions"),
        ("Compliance Officer",      "AML Alert referrals, compliance breach escalations"),
    ]
    table = doc.add_table(rows=1 + len(roles), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    set_cell_bg(hdr.cells[0], BRAND)
    set_cell_bg(hdr.cells[1], BRAND)
    cell_para(hdr.cells[0], "Role",             bold=True, color=WHITE, size=9)
    cell_para(hdr.cells[1], "Responsibilities", bold=True, color=WHITE, size=9)
    for i, (role, resp) in enumerate(roles):
        r = table.rows[i + 1]
        bg = RGBColor(0xF0, 0xFA, 0xFA) if i % 2 == 0 else WHITE
        set_cell_bg(r.cells[0], bg)
        set_cell_bg(r.cells[1], bg)
        cell_para(r.cells[0], role, bold=True, size=9, color=BRAND_DARK)
        cell_para(r.cells[1], resp, size=9)
    doc.add_paragraph()

    # ── 5. Systems Referenced
    heading1(doc, "5.  Systems Referenced")
    systems = [
        ("LoanIQ",                     "Primary loan management and origination system"),
        ("Compliance Tracker",          "AML screening, PEP checks, adverse media review"),
        ("Document Management System",  "Document storage, completeness certificates, commitment letters"),
        ("Loan Application Worksheet",  "DTI calculation, credit certification"),
        ("Appraisal Management System", "Appraisal vendor engagement and tracking"),
        ("Credit Committee System",     "Sanctioning workflow and decision recording"),
    ]
    table = doc.add_table(rows=1 + len(systems), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    set_cell_bg(hdr.cells[0], BRAND)
    set_cell_bg(hdr.cells[1], BRAND)
    cell_para(hdr.cells[0], "System",  bold=True, color=WHITE, size=9)
    cell_para(hdr.cells[1], "Purpose", bold=True, color=WHITE, size=9)
    for i, (sys, pur) in enumerate(systems):
        r = table.rows[i + 1]
        bg = RGBColor(0xF0, 0xFA, 0xFA) if i % 2 == 0 else WHITE
        set_cell_bg(r.cells[0], bg)
        set_cell_bg(r.cells[1], bg)
        cell_para(r.cells[0], sys, bold=True, size=9, color=BRAND_DARK)
        cell_para(r.cells[1], pur, size=9)
    doc.add_paragraph()

    # ── 6. Key Policies
    heading1(doc, "6.  Key Policies and Compliance References")
    policies = [
        ("POL-AML-001",  "Anti-Money Laundering Policy",    "AML screen mandatory before credit bureau query (Section 2.3)"),
        ("POL-DOC-002",  "Document Completeness Policy",    "Certificate required before appraisal instruction (Section 3.1)"),
        ("POL-CRED-003", "DTI Certification Policy",        "Credit Analyst sign-off only (Section 4.1)"),
        ("POL-SANC-004", "Sanctioning Timeline Policy",     "1 business day SLA (Section 5.2)"),
        ("POL-COMP-005", "Document Upload SLA Policy",      "80% of applications within 48 hours (Section 2.1)"),
    ]
    table = doc.add_table(rows=1 + len(policies), cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for c, lbl in enumerate(("Policy Ref", "Policy Name", "Requirement")):
        set_cell_bg(hdr.cells[c], BRAND)
        cell_para(hdr.cells[c], lbl, bold=True, color=WHITE, size=9)
    for i, (ref, name, req) in enumerate(policies):
        r = table.rows[i + 1]
        bg = RGBColor(0xF0, 0xFA, 0xFA) if i % 2 == 0 else WHITE
        for c in range(3):
            set_cell_bg(r.cells[c], bg)
        cell_para(r.cells[0], ref,  bold=True, size=9, color=BRAND_DARK)
        cell_para(r.cells[1], name, size=9)
        cell_para(r.cells[2], req,  size=9)
    doc.add_paragraph()

    # ── 7. SLA Summary
    heading1(doc, "7.  SLA Summary")
    slas = [
        ("AML Pre-Screening",         "Loan Processing Officer",           "4 business hours"),
        ("Credit Bureau Pull",         "Loan Processing Officer",           "1 business day"),
        ("Document Upload",            "Loan Processing Officer",           "48 hours (80% threshold)"),
        ("DTI Calculation",            "Credit Analyst",                    "1 business day"),
        ("Property Appraisal Report",  "Appraisal Vendor",                  "5 business days"),
        ("Credit Assessment",          "Credit Analyst / Senior Underwriter","2 business days"),
        ("Sanctioning Decision",       "Credit Committee",                  "1 business day"),
        ("Commitment Letter",          "Branch Manager",                    "1 business day"),
    ]
    table = doc.add_table(rows=1 + len(slas), cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for c, lbl in enumerate(("Stage / Activity", "Responsible Role", "SLA")):
        set_cell_bg(hdr.cells[c], BRAND)
        cell_para(hdr.cells[c], lbl, bold=True, color=WHITE, size=9)
    for i, (act, role, sla) in enumerate(slas):
        r = table.rows[i + 1]
        bg = RGBColor(0xF0, 0xFA, 0xFA) if i % 2 == 0 else WHITE
        for c in range(3):
            set_cell_bg(r.cells[c], bg)
        cell_para(r.cells[0], act,  bold=True, size=9)
        cell_para(r.cells[1], role, size=9)
        cell_para(r.cells[2], sla,  size=9)
    doc.add_paragraph()

    # ── Footer approval block
    doc.add_paragraph()
    approval_data = [
        ("Approved by",  "Head of Commercial Banking"),
        ("Reviewed by",  "Chief Risk Officer, Head of Compliance"),
        ("Next Review",  "1 March 2025"),
    ]
    table = doc.add_table(rows=len(approval_data), cols=2)
    table.style = 'Table Grid'
    for i, (lbl, val) in enumerate(approval_data):
        set_cell_bg(table.rows[i].cells[0], RGBColor(0xE6, 0xF4, 0xF4))
        cell_para(table.rows[i].cells[0], lbl, bold=True, size=9, color=BRAND_DARK)
        cell_para(table.rows[i].cells[1], val, size=9)

    path = OUT / "commercial_banking_loan_sop.docx"
    doc.save(path)
    print(f"Saved: {path}")


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 2 — AUDIT REPORT
# ══════════════════════════════════════════════════════════════════════

def build_audit_report():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    add_header_banner(
        doc,
        "COMMERCIAL BANKING DIVISION — INTERNAL AUDIT",
        "Loan Origination Quarterly Compliance Audit Report — Q3 2024",
        "IA-2024-Q3-LO-007",
        "Review Period: 1 Jul – 30 Sep 2024",
        "Issued October 2024",
        "Internal Audit Department",
    )

    # ── 1. Executive Summary
    heading1(doc, "1.  Executive Summary")
    body(doc,
        "This report presents the findings of the Internal Audit Department's quarterly review "
        "of loan origination processes within the Commercial Banking Division. The review covered "
        "12 loan applications processed between 1 July and 30 September 2024 across three branch "
        "locations (Downtown, Riverside, and Eastgate)."
    )
    body(doc,
        "Overall process adherence was found to be partial. While credit assessment activities and "
        "post-approval documentation were generally compliant, significant deviations from SOP v4.2 "
        "(effective March 2024) were identified in AML pre-screening sequencing, appraisal ordering "
        "protocols, sanctioning timelines, and role-specific sign-off requirements."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Management attention is drawn to the AML sequencing finding, which represents a "
                  "regulatory compliance risk and requires immediate remediation.")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RED
    doc.add_paragraph()

    # ── 2. Scope and Methodology
    heading1(doc, "2.  Scope and Methodology")
    body(doc,
        "The audit examined loan origination files for 12 applications across the Downtown, "
        "Riverside, and Eastgate branch offices. Auditors reviewed loan origination documentation, "
        "system access logs from the LoanIQ platform, timestamps recorded in the Document Management "
        "System, Compliance Tracker entries, and approval records from the Credit Committee system."
    )
    body(doc,
        "Each loan file was assessed against the current Loan Origination SOP version 4.2. Where "
        "timestamps were unavailable, auditors relied on document dating and system-generated activity records."
    )

    # ── 3. Findings
    heading1(doc, "3.  Findings")

    findings = [
        {
            "title":  "3.1  Credit Bureau Pull",
            "status": "COMPLIANT",
            "color":  GREEN,
            "body": [
                "Credit bureau enquiries (Experian and TransUnion dual pull) were completed and formally "
                "documented for all 12 reviewed applications prior to the commencement of credit analysis. "
                "No exceptions were noted. Bureau reports were filed within the designated folder structure "
                "in the Document Management System in all cases.",
            ],
        },
        {
            "title":  "3.2  Debt-to-Income Ratio Calculation",
            "status": "GENERALLY COMPLIANT",
            "color":  RGBColor(0xD9, 0x77, 0x06),
            "body": [
                "DTI ratio calculations were present in all 12 reviewed files and formally recorded on "
                "the standard Loan Application Worksheet. In 11 of 12 cases, the calculation was performed "
                "and certified by a Credit Analyst as required under SOP Section 4.1.",
                "However, in one case reviewed at the Eastgate branch, the DTI calculation was performed "
                "and certified by a Loan Officer rather than a Credit Analyst. The calculation result was "
                "arithmetically correct, but the sign-off authority was not in accordance with the role "
                "requirements prescribed in the SOP. This constitutes a role deviation that must be "
                "corrected in future applications.",
            ],
        },
        {
            "title":  "3.3  Document Upload and Completeness",
            "status": "COMPLIANT",
            "color":  GREEN,
            "body": [
                "Loan Processing Officers uploaded required customer documentation within the 48-hour SLA "
                "in 10 of 12 reviewed cases (83%), which meets the acceptable compliance threshold of 80% "
                "set out in SOP Section 2.1. In the two non-compliant cases, delays ranged from 52 to 61 "
                "hours and were attributable to late submission by applicants.",
            ],
        },
        {
            "title":  "3.4  AML and Compliance Pre-Screening",
            "status": "NON-COMPLIANT (CRITICAL)",
            "color":  RED,
            "body": [
                "Anti-Money Laundering screening was not completed before the credit bureau pull was "
                "initiated in 4 of the 12 reviewed cases (33%). Compliance Tracker system logs confirm "
                "that the AML check was recorded after the bureau query had already been submitted in "
                "these instances.",
                "This ordering directly reverses the mandatory sequence prescribed in SOP Section 2.3, "
                "which states that AML screening must be completed and cleared as a prerequisite gate "
                "before any credit enquiry is made. This finding has regulatory implications and is "
                "classified as a critical deviation.",
            ],
        },
        {
            "title":  "3.5  Property Appraisal Sequencing",
            "status": "NON-COMPLIANT",
            "color":  RED,
            "body": [
                "In 2 of the 12 reviewed cases, the property appraisal was commissioned before document "
                "completeness had been formally confirmed. SOP Section 3.1 requires that a Document "
                "Completeness Certificate be issued and countersigned before the appraisal vendor is engaged.",
                "In both instances, the appraisal was ordered directly by the Loan Processing Officer "
                "without the required completeness sign-off. One appraisal had to be paused mid-instruction "
                "when the missing documentation was subsequently identified.",
            ],
        },
        {
            "title":  "3.6  Sanctioning Timeline",
            "status": "NON-COMPLIANT",
            "color":  RED,
            "body": [
                "Sanctioning decisions were issued within the one business day SLA set out in SOP Section "
                "5.2 in only 5 of the 12 reviewed cases. The average sanctioning timeline across all "
                "reviewed files was 3.2 business days.",
                "In 3 cases, delays exceeded 5 business days with no documented escalation to the Credit "
                "Committee. The longest delay recorded was 8 business days for a commercial property "
                "application at the Riverside branch.",
            ],
        },
        {
            "title":  "3.7  Commitment Letter Issuance",
            "status": "COMPLIANT",
            "color":  GREEN,
            "body": [
                "Commitment letters were issued to customers following Credit Committee approval in all "
                "12 reviewed cases where approval was granted. Letter content, fee disclosures, and "
                "signatory authority were reviewed and found to be compliant with current policy "
                "requirements in all instances.",
            ],
        },
    ]

    for f in findings:
        heading2(doc, f["title"])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Status: ")
        r.bold = True
        r.font.size = Pt(9)
        status_badge(p, f["status"], f["color"])
        for para in f["body"]:
            body(doc, para)

    # ── 4. Recommendations
    heading1(doc, "4.  Recommendations")

    recs = [
        ("4.1", "AML Pre-Screening Gate",            "IMMEDIATE",
         "Implement a system-level workflow gate in LoanIQ to prevent credit bureau queries from being "
         "submitted until the AML Compliance Tracker record shows a Cleared status. Conduct mandatory "
         "retraining for all Loan Processing Officers on the AML sequencing requirement within 30 days."),
        ("4.2", "Document Completeness Before Appraisal", "HIGH",
         "Enforce the Document Completeness Certificate as a mandatory system dependency before the "
         "appraisal order workflow step can be initiated. Review the two affected cases to determine "
         "whether additional steps are required."),
        ("4.3", "Sanctioning Timeline Monitoring",   "HIGH",
         "Branch managers are required to review sanctioning queue status daily and initiate escalation "
         "to the Credit Committee for any application exceeding 2 business days without a decision. "
         "A management dashboard report should be configured in LoanIQ to surface breach cases automatically."),
        ("4.4", "Role-Specific Sign-Off Controls",   "MEDIUM",
         "Review access permissions in the Loan Application Worksheet to ensure that the DTI certification "
         "field can only be completed by users with the Credit Analyst role designation."),
    ]

    priority_colors = {
        "IMMEDIATE": RED,
        "HIGH":      RGBColor(0xD9, 0x77, 0x06),
        "MEDIUM":    RGBColor(0x0E, 0x79, 0xB2),
    }

    for ref, title, priority, text in recs:
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        set_cell_bg(cell, RGBColor(0xF9, 0xFB, 0xFB))
        p = cell.paragraphs[0]
        p.clear()
        r1 = p.add_run(f"{ref}  {title}    ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = BRAND_DARK
        r2 = p.add_run(priority)
        r2.bold = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = priority_colors.get(priority, GREY_MID)
        p2 = cell.add_paragraph(text)
        p2.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # ── 5. Management Response
    heading1(doc, "5.  Management Response")
    body(doc,
        "Management accepts all four recommendations. Remediation plans for items 4.1 and 4.4 are "
        "targeted for completion within 45 days. Items 4.2 and 4.3 are targeted for the next system "
        "release cycle."
    )

    # ── Signatories
    doc.add_paragraph()
    sig_data = [
        ("Prepared by",   "Senior Internal Auditor, Commercial Banking"),
        ("Reviewed by",   "Head of Internal Audit"),
        ("Distribution",  "Chief Risk Officer, Head of Commercial Banking, Branch Managers (Downtown, Riverside, Eastgate)"),
    ]
    table = doc.add_table(rows=len(sig_data), cols=2)
    table.style = 'Table Grid'
    for i, (lbl, val) in enumerate(sig_data):
        set_cell_bg(table.rows[i].cells[0], RGBColor(0xE6, 0xF4, 0xF4))
        cell_para(table.rows[i].cells[0], lbl, bold=True, size=9, color=BRAND_DARK)
        cell_para(table.rows[i].cells[1], val, size=9)

    path = OUT / "loan_audit_report.docx"
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build_sop()
    build_audit_report()
    print("Done.")
