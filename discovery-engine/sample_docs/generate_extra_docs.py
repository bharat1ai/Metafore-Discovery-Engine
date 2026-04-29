"""Generate three additional banking demo documents for testing multi-document upload.

Produces:
  - compliance_policy.docx       (~600 words, AML/KYC + fair lending — omits LOS refs;
                                   Compliance Officer role has no process steps)
  - role_descriptions.docx       (~400 words, 5 roles — includes a Quality Assurance Officer
                                   that does not appear in any SOP; Credit Analyst reports to
                                   Branch Credit Manager — contradicts SOP context)
  - it_systems_register.docx     (~400 words, IT systems — includes Document Management
                                   Portal that appears in no SOP; LOS integrates with
                                   AppraisalTrack — contradicts SOP TBD note)

These deliberate gaps and contradictions surface as Cross-Document Insights when all
three are uploaded alongside commercial_banking_loan_sop.docx.

Reuses the styling helpers from generate_docx.py for visual consistency.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Reuse helpers + colour tokens from the existing generator.
from generate_docx import (
    BRAND, BRAND_DARK, WHITE, RED, GREEN,
    add_header_banner, heading1, heading2, body,
    meta_table, set_cell_bg, cell_para,
)

OUT = Path(__file__).parent


# ══════════════════════════════════════════════════════════════════════
# DOC 1 — Compliance Policy (~600 words)
# ══════════════════════════════════════════════════════════════════════

def build_compliance_policy():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    add_header_banner(
        doc,
        "RETAIL BANKING — FINANCIAL CRIME COMPLIANCE",
        "Anti-Money Laundering, Customer Due Diligence and Fair Lending Policy",
        "POL-COMP-2024-002",
        "Version 2.1",
        "1 February 2024",
        "Head of Compliance and Financial Crime",
    )

    # 1. Purpose and Scope
    heading1(doc, "1.  Purpose and Scope")
    body(doc,
        "This Policy sets out the minimum control standards required to prevent the bank from being "
        "used to launder the proceeds of crime, finance terrorism, breach economic sanctions, or "
        "engage in unlawful discrimination in the extension of credit. It applies to all retail "
        "lending and account-opening activity across all branches, channels and product lines."
    )
    body(doc,
        "Adherence is mandatory. Suspected breaches must be reported to the Compliance Officer "
        "without delay and, where applicable, escalated under the Financial Crime Reporting Standard."
    )

    # 2. AML and Customer Due Diligence
    heading1(doc, "2.  Anti-Money Laundering and Customer Due Diligence")

    heading2(doc, "2.1  Customer Identification and Verification")
    body(doc,
        "All applicants must be subject to Customer Due Diligence prior to the establishment of any "
        "credit relationship. CDD comprises identity verification, beneficial ownership disclosure for "
        "non-individual applicants, source-of-funds enquiry, and an initial risk rating. Enhanced Due "
        "Diligence applies to higher-risk jurisdictions, complex ownership structures, and PEP relationships."
    )

    heading2(doc, "2.2  Sanctions, PEP and Adverse Media Screening")
    body(doc,
        "Every applicant, beneficial owner, and authorised signatory must be screened against current "
        "sanctions lists (OFAC, UN, EU, HMT), recognised PEP databases, and adverse media sources. "
        "Positive matches must be referred to the Compliance Officer within four hours and may not be "
        "auto-cleared by front-office personnel."
    )

    heading2(doc, "2.3  Sequencing Requirement (Mandatory)")
    p = doc.add_paragraph()
    r = p.add_run(
        "AML and sanctions screening MUST be completed and a Cleared status formally recorded BEFORE "
        "any credit bureau enquiry is initiated for the applicant. Conducting a credit bureau enquiry "
        "while AML screening is open or in Alert status is a reportable policy breach. This sequence "
        "is non-negotiable and supersedes any operational SOP that may suggest otherwise."
    )
    r.font.size = Pt(9.5); r.bold = True; r.font.color.rgb = RED

    heading2(doc, "2.4  Suspicious Activity Reporting")
    body(doc,
        "Where activity inconsistent with the customer profile is identified, a Suspicious Activity "
        "Report must be raised through the Financial Crime Case Management standard. Tipping-off is "
        "strictly prohibited."
    )

    # 3. Fair Lending
    heading1(doc, "3.  Fair Lending and Equal Credit Opportunity")
    body(doc,
        "Credit decisions must be made solely on demonstrable creditworthiness. Discrimination on the "
        "basis of race, color, religion, national origin, sex, marital status, age, or receipt of public "
        "assistance income is strictly prohibited under the Equal Credit Opportunity Act (Regulation B) "
        "and the Fair Housing Act."
    )
    body(doc,
        "All adverse action notices must be issued within 30 days of the credit decision and must state "
        "the specific reasons for denial. Disparate impact analysis is performed quarterly by the "
        "Compliance team."
    )

    # 4. Roles and Accountabilities
    heading1(doc, "4.  Roles and Accountabilities")
    roles_data = [
        ("Compliance Officer",
         "Owns this Policy. Reviews escalations from front-office. Approves SARs. Conducts thematic compliance reviews. "
         "The Compliance Officer is an oversight role and does NOT execute application-processing steps; ownership of "
         "operational sequencing rests with the lines of business under their respective SOPs."),
        ("Branch Manager",
         "Accountable for first-line adherence to this Policy within the branch. Required to confirm adherence on a "
         "quarterly attestation."),
        ("All Front-Office Staff",
         "Required to complete this Policy's mandatory training annually and to refer any sanctions, PEP, or adverse "
         "media match to the Compliance Officer without independent action."),
    ]
    table = doc.add_table(rows=1 + len(roles_data), cols=2)
    table.style = 'Table Grid'
    set_cell_bg(table.rows[0].cells[0], BRAND); set_cell_bg(table.rows[0].cells[1], BRAND)
    cell_para(table.rows[0].cells[0], "Role", bold=True, color=WHITE, size=9)
    cell_para(table.rows[0].cells[1], "Accountability", bold=True, color=WHITE, size=9)
    for i, (rl, ac) in enumerate(roles_data):
        r = table.rows[i + 1]
        bg = RGBColor(0xF0, 0xFA, 0xFA) if i % 2 == 0 else WHITE
        set_cell_bg(r.cells[0], bg); set_cell_bg(r.cells[1], bg)
        cell_para(r.cells[0], rl, bold=True, size=9, color=BRAND_DARK)
        cell_para(r.cells[1], ac, size=9)
    doc.add_paragraph()

    # 5. Reporting Lines and Escalation
    heading1(doc, "5.  Reporting and Escalation")
    body(doc,
        "All Compliance Officer determinations are reported to the Chief Compliance Officer monthly and to the "
        "Audit Committee quarterly. Material incidents must be escalated to the Chief Risk Officer within 24 hours."
    )

    # 6. Policy Cross-References
    heading1(doc, "6.  Cross-Referenced Policies")
    body(doc,
        "POL-AML-001 Anti-Money Laundering Standard;  POL-DOC-002 Document Completeness Policy;  "
        "POL-COMP-005 Document Upload SLA Policy;  POL-FAIR-003 Fair Lending Standard;  "
        "POL-SANC-006 Sanctions Screening Standard."
    )

    # Sign-off block
    doc.add_paragraph()
    sig = [("Approved by", "Chief Compliance Officer"),
           ("Reviewed by", "General Counsel; Chief Risk Officer"),
           ("Next Review", "1 February 2025")]
    table = doc.add_table(rows=len(sig), cols=2); table.style = 'Table Grid'
    for i, (lbl, val) in enumerate(sig):
        set_cell_bg(table.rows[i].cells[0], RGBColor(0xE6, 0xF4, 0xF4))
        cell_para(table.rows[i].cells[0], lbl, bold=True, size=9, color=BRAND_DARK)
        cell_para(table.rows[i].cells[1], val, size=9)

    path = OUT / "compliance_policy.docx"
    doc.save(path); print(f"Saved: {path}")


# ══════════════════════════════════════════════════════════════════════
# DOC 2 — Role Descriptions (~400 words)
# ══════════════════════════════════════════════════════════════════════

def build_role_descriptions():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    add_header_banner(
        doc,
        "HUMAN RESOURCES — COMMERCIAL BANKING",
        "Loan Origination Department — Role Descriptions",
        "HR-RD-2024-LO-V1",
        "Version 1.0",
        "March 2024",
        "HR Business Partner, Commercial Banking",
    )

    heading1(doc, "1.  Purpose")
    body(doc,
        "This document sets out the formal role descriptions for the five operational positions within "
        "the Loan Origination Department. It is the authoritative reference for hiring, performance "
        "management, and reporting-line decisions. Where this document and an operational SOP differ, "
        "this document governs reporting-line and accountability matters."
    )

    heading1(doc, "2.  Roles")

    # Role 1
    heading2(doc, "2.1  Loan Processing Officer")
    body(doc,
        "Front-office role responsible for intake of new loan applications and the operational steps that "
        "follow customer engagement. Reports to the Branch Manager. Performs document collection, system "
        "data entry, AML screening initiation, and bureau enquiry preparation. Minimum qualification: "
        "Associate degree plus two years of branch operations experience."
    )

    # Role 2 — contradicting reporting line
    heading2(doc, "2.2  Credit Analyst")
    p = doc.add_paragraph()
    r = p.add_run(
        "Credit Analysts are responsible for the certification of Debt-to-Income calculations and the "
        "preparation of credit assessment reports. Credit Analysts report to the Branch Credit Manager "
        "and are administratively aligned to their home branch."
    )
    r.font.size = Pt(9.5)
    body(doc,
        "Credit Analysts hold the certifying authority for DTI calculations within the Loan Application "
        "Worksheet. Loan Processing Officers do not have authority to certify DTI calculations. Minimum "
        "qualification: Bachelor's degree in finance, accounting or related field, plus the bank's "
        "internal Credit Analyst certification."
    )

    # Role 3
    heading2(doc, "2.3  Senior Underwriter")
    body(doc,
        "Senior Underwriters take responsibility for credit assessment of higher-risk applications, "
        "specifically those with Debt-to-Income ratios exceeding 43%. Reports to the Head of Underwriting. "
        "Minimum qualification: Five years' commercial credit experience plus the bank's Senior Underwriter "
        "accreditation."
    )

    # Role 4
    heading2(doc, "2.4  Branch Manager")
    body(doc,
        "Accountable for branch-level performance and adherence to all procedural standards. Counter-signs "
        "Document Completeness Certificates, monitors sanctioning queue status, and is the sole signatory "
        "for Commitment Letters issued by the branch. Reports to the Regional Director, Commercial Banking."
    )

    # Role 5 — orphan, appears in no SOP
    heading2(doc, "2.5  Quality Assurance Officer")
    body(doc,
        "The Quality Assurance Officer is an independent second-line role attached to each branch. "
        "Conducts post-decision sample reviews of completed loan files, tracks defect trends, and reports "
        "monthly findings to the Branch Manager and the Regional Director. The Quality Assurance Officer "
        "does NOT participate in front-line application processing and is not referenced in operational "
        "SOPs governing day-to-day origination activity. Reports to the Regional Director, Commercial "
        "Banking. Minimum qualification: Bachelor's degree plus three years' audit or quality assurance "
        "experience."
    )

    # Footer
    doc.add_paragraph()
    sig = [("Approved by", "Head of HR, Commercial Banking"),
           ("Effective", "March 2024"),
           ("Review Cycle", "Annual")]
    table = doc.add_table(rows=len(sig), cols=2); table.style = 'Table Grid'
    for i, (lbl, val) in enumerate(sig):
        set_cell_bg(table.rows[i].cells[0], RGBColor(0xE6, 0xF4, 0xF4))
        cell_para(table.rows[i].cells[0], lbl, bold=True, size=9, color=BRAND_DARK)
        cell_para(table.rows[i].cells[1], val, size=9)

    path = OUT / "role_descriptions.docx"
    doc.save(path); print(f"Saved: {path}")


# ══════════════════════════════════════════════════════════════════════
# DOC 3 — IT Systems Register (~400 words)
# ══════════════════════════════════════════════════════════════════════

def build_it_systems_register():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    add_header_banner(
        doc,
        "IT SERVICE MANAGEMENT — COMMERCIAL BANKING",
        "Loan Department IT Systems Register",
        "IT-SR-2024-LO-Q1",
        "Q1 2024 Edition",
        "April 2024",
        "IT Service Management Office",
    )

    heading1(doc, "1.  Purpose")
    body(doc,
        "This Register is the authoritative inventory of IT systems used by the Loan Origination "
        "Department. It records system ownership, integration relationships, data classification, and "
        "compliance status. The Register is the source of truth for system-of-record disputes between "
        "operational SOPs and integration documentation."
    )

    heading1(doc, "2.  Systems Inventory")

    systems = [
        ("LoanIQ (LOS)",
         "Origination",
         "Integrates with AppraisalTrack, Compliance Tracker, and Credit Bureau Gateway. Note: this "
         "Register confirms a live, automated LoanIQ ↔ AppraisalTrack integration despite an outstanding "
         "TBD note in SOP-LO-004.2 Section 3.5. The integration status was certified by IT Operations on "
         "12 January 2024 and supersedes the SOP entry."),
        ("Compliance Tracker",
         "Financial Crime",
         "AML, sanctions, PEP and adverse media screening. Owned by Financial Crime Compliance. "
         "Integrates with LoanIQ for Cleared status confirmation."),
        ("Credit Bureau Gateway",
         "Credit",
         "API-based dual bureau service (Experian, TransUnion). Invoked by LoanIQ; results filed "
         "automatically in the Document Management System."),
        ("AppraisalTrack",
         "Appraisal",
         "Vendor-managed appraisal instruction and reporting platform. Replaces the legacy Appraisal "
         "Management System referenced in older SOPs."),
        ("Document Management System (DMS)",
         "Records",
         "Bank-wide document repository. Holds bureau reports, completeness certificates, commitment "
         "letters, appraisal reports, and credit assessment files."),
        ("Document Management Portal (DMP)",
         "Self-Service",
         "Customer-facing document submission portal recently commissioned by the Digital Channels team. "
         "Allows applicants to upload supporting documentation directly. Not currently referenced in any "
         "Loan Origination SOP. Pending review by the Loan Origination process governance group for "
         "potential integration into the document collection step."),
        ("Credit Committee Workflow System",
         "Decisioning",
         "Workflow platform supporting credit committee deliberation, voting, and decision recording. "
         "Integrated with LoanIQ for sanctioning status updates."),
    ]
    table = doc.add_table(rows=1 + len(systems), cols=3); table.style = 'Table Grid'
    hdr = table.rows[0]
    for c, lbl in enumerate(("System", "Domain", "Description and Integrations")):
        set_cell_bg(hdr.cells[c], BRAND)
        cell_para(hdr.cells[c], lbl, bold=True, color=WHITE, size=9)
    for i, (sys_name, dom, desc) in enumerate(systems):
        r = table.rows[i + 1]
        bg = RGBColor(0xF0, 0xFA, 0xFA) if i % 2 == 0 else WHITE
        for c in range(3):
            set_cell_bg(r.cells[c], bg)
        cell_para(r.cells[0], sys_name, bold=True, size=9, color=BRAND_DARK)
        cell_para(r.cells[1], dom, size=9)
        cell_para(r.cells[2], desc, size=9)
    doc.add_paragraph()

    heading1(doc, "3.  Change Control")
    body(doc,
        "Any change to a system listed above — including new integrations, ownership transfer, or "
        "decommissioning — must be raised through the IT Change Advisory Board. Operational SOPs must "
        "be reviewed and updated to align with this Register following any approved change."
    )

    # Footer
    doc.add_paragraph()
    sig = [("Maintained by", "IT Service Management Office"),
           ("Last verified", "April 2024"),
           ("Next refresh", "Quarterly")]
    table = doc.add_table(rows=len(sig), cols=2); table.style = 'Table Grid'
    for i, (lbl, val) in enumerate(sig):
        set_cell_bg(table.rows[i].cells[0], RGBColor(0xE6, 0xF4, 0xF4))
        cell_para(table.rows[i].cells[0], lbl, bold=True, size=9, color=BRAND_DARK)
        cell_para(table.rows[i].cells[1], val, size=9)

    path = OUT / "it_systems_register.docx"
    doc.save(path); print(f"Saved: {path}")


if __name__ == "__main__":
    build_compliance_policy()
    build_role_descriptions()
    build_it_systems_register()
    print("Done.")
