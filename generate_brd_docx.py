"""Generate a professional DOCX version of the Discovery Engine BRD."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent

BRAND      = RGBColor(0x03, 0x68, 0x68)
BRAND_DARK = RGBColor(0x02, 0x4F, 0x4F)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TEAL_LIGHT = RGBColor(0xE6, 0xF4, 0xF4)
TEAL_MID   = RGBColor(0xA7, 0xD4, 0xD4)
TEXT_DARK  = RGBColor(0x1F, 0x2D, 0x2D)
TEXT_MID   = RGBColor(0x4A, 0x68, 0x68)
GREY       = RGBColor(0x6B, 0x72, 0x80)
AMBER      = RGBColor(0xD9, 0x77, 0x06)
GREEN      = RGBColor(0x16, 0xA3, 0x4A)
RED        = RGBColor(0xDC, 0x26, 0x26)
GOLD       = RGBColor(0xB4, 0x5A, 0x00)


# ── Helpers ───────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb):
    hex_color = format(rgb[0], '02X') + format(rgb[1], '02X') + format(rgb[2], '02X') \
        if isinstance(rgb, tuple) else str(rgb)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def cell_para(cell, text, bold=False, size=9, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.paragraphs[0].alignment = align
    return cell.paragraphs[0]


def add_run(para, text, bold=False, italic=False, size=None, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def section_heading(doc, number, title):
    """Numbered section heading with teal top border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, f"{number}  ", bold=True, size=13, color=BRAND)
    add_run(p, title.upper(), bold=True, size=11, color=BRAND_DARK)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '6')
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), '036868')
    pBdr.append(top)
    pPr.append(pBdr)


def sub_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, title, bold=True, size=10, color=BRAND_DARK)


def body_text(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXT_DARK


def bullet(doc, text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + indent * 0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_DARK


def add_table(doc, headers, rows, col_widths=None, alt_rows=True):
    """Generic branded table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for c, h in enumerate(headers):
        set_cell_bg(table.rows[0].cells[c], BRAND)
        cell_para(table.rows[0].cells[c], h, bold=True, color=WHITE, size=9)
    for i, row in enumerate(rows):
        bg = TEAL_LIGHT if (alt_rows and i % 2 == 0) else WHITE
        for c, val in enumerate(row):
            set_cell_bg(table.rows[i + 1].cells[c], bg)
            cell_para(table.rows[i + 1].cells[c], str(val), size=9)
    if col_widths:
        for r in table.rows:
            for c, w in enumerate(col_widths):
                r.cells[c].width = Inches(w)
    doc.add_paragraph()
    return table


def add_note_box(doc, text, color=TEAL_LIGHT, border_color='036868'):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_DARK
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# SECTIONS
# ══════════════════════════════════════════════════════════════════════

def add_cover(doc):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, BRAND_DARK)

    p1 = cell.paragraphs[0]; p1.clear(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("METAFORE WORKS")
    r.font.color.rgb = TEAL_MID; r.font.size = Pt(9); r.font.bold = True

    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Discovery Engine")
    r2.font.color.rgb = WHITE; r2.font.size = Pt(24); r2.font.bold = True

    p3 = cell.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Business Requirements Document")
    r3.font.color.rgb = TEAL_MID; r3.font.size = Pt(13)

    p4 = cell.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Version 1.0   |   April 2026   |   CONFIDENTIAL")
    r4.font.color.rgb = TEAL_MID; r4.font.size = Pt(8.5)

    doc.add_paragraph()


def add_meta_table(doc):
    rows = [
        ("Document Title",  "Metafore Works Discovery Engine — Business Requirements Document"),
        ("Version",         "1.0"),
        ("Date",            "April 2026"),
        ("Status",          "Active"),
        ("Author",          "Metafore Works"),
        ("Classification",  "Confidential"),
    ]
    add_table(doc, ["Field", "Detail"], rows, col_widths=[1.5, 5.0])


def add_s1_summary(doc):
    section_heading(doc, "1", "Executive Summary")
    body_text(doc,
        "The Metafore Works Discovery Engine is an AI-powered document intelligence platform "
        "that transforms unstructured business documents — SOPs, policies, audit reports — into "
        "structured knowledge graphs. It enables organisations to visualise process flows, identify "
        "compliance gaps, generate domain object models, and assess conformance against operating "
        "standards, all through a browser-based interface with no installation required beyond Python."
    )
    add_note_box(doc,
        "The Discovery Engine is the foundational layer of the Metafore Works platform. "
        "It is designed to be the first step before any application build — understanding "
        "the organisation before designing for it."
    )


def add_s2_context(doc):
    section_heading(doc, "2", "Business Context and Problem Statement")
    body_text(doc,
        "Organisations accumulate large volumes of operational documentation containing rich process "
        "and governance knowledge locked in unstructured text. Extracting, relating, and acting on "
        "this knowledge currently requires significant manual analyst effort."
    )
    add_table(doc,
        ["Problem", "Impact"],
        [
            ("Process knowledge buried in long-form documents",      "Analysts spend days manually mapping workflows"),
            ("Gap analysis done manually and inconsistently",        "Critical compliance gaps are missed or under-reported"),
            ("Object models created from scratch each time",         "Data model rework delays system design phases"),
            ("No automated conformance checking against SOPs",       "Audit findings remain disconnected from source procedures"),
        ],
        col_widths=[3.0, 3.5]
    )


def add_s3_objectives(doc):
    section_heading(doc, "3", "Objectives")
    objectives = [
        "Extract a structured knowledge graph from any uploaded business document in under 60 seconds.",
        "Mine the live operational data set against the extracted SOP via a Celonis-style Process Mining surface — fitness, bottleneck, deviation patterns, top deviating cases — with no LLM dependency on the headline view.",
        "Generate AI optimisation suggestions (re-routing, SLA targets, role assignments) grounded in the Process Mining snapshot.",
        "Automatically generate a BRD §4-compliant domain object model (Pydantic + JSON Schema + Entity Diagram).",
        "Calculate gap analysis and provide actionable remediation blueprints.",
        "Score Audit Check (was Conformance) between an audit document and the source process graph; disambiguated from Process Mining → Execution Conformance.",
        "Provide natural language querying over the knowledge graph.",
        "Produce a print-ready Executive Report consolidating every feature run, with Metafore brand typography and page breaks.",
    ]
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_run(p, f"{i}.  ", bold=True, size=9.5, color=BRAND)
        add_run(p, obj, size=9.5, color=TEXT_DARK)


def add_s4_stakeholders(doc):
    section_heading(doc, "4", "Stakeholders")
    add_table(doc,
        ["Role", "Interest"],
        [
            ("Product Owner — Metafore Works",  "Feature prioritisation and roadmap"),
            ("Business Analyst",                "Gap analysis, audit check, object model output"),
            ("Data Architect",                  "BRD §4 object model export and ERD review"),
            ("Process Owner",                   "Process Mining bottleneck / deviations and AI Optimisation suggestions"),
            ("Compliance / Audit",              "Audit Check (doc-vs-doc) and Execution Conformance (data-vs-doc)"),
            ("IT / Infrastructure",             "Deployment and API key management"),
        ],
        col_widths=[2.5, 4.0]
    )


def add_s5_scope(doc):
    section_heading(doc, "5", "Scope")
    sub_heading(doc, "5.1  In Scope")
    in_scope = [
        "Document upload: PDF, DOCX, TXT",
        "AI-driven knowledge graph extraction (nodes and edges)",
        "Interactive graph visualisation via vis-network",
        "Process Mining (Workflows tab) — Process Map, Variants, Execution Conformance, Filter Cases, AI Optimise (Haiku)",
        "Gap analysis with remediation blueprint",
        "Pulse drawer with severity filter, Source line, Dismiss action, AI Strategic Insights",
        "Audit Check (renamed from standalone Conformance) — document-only, doc-vs-doc via Sonnet",
        "Object model generation: Entity Diagram (default tab) + Pydantic v2 + JSON Schema, with cache-busting Regenerate",
        "Natural language querying over the graph (grounded in graph + live operational data)",
        "Dashboard — full-screen, hotspots and deviating cases sourced from Process Mining",
        "Executive Report — print-ready HTML with Metafore brand typography, page breaks per section",
        "Single-user, local deployment on Windows / Mac",
    ]
    for item in in_scope:
        bullet(doc, item)
    doc.add_paragraph()

    sub_heading(doc, "5.2  Out of Scope")
    out_scope = [
        "Multi-user or multi-tenant access",
        "Persistent database storage (all data is in-memory, reset on server restart)",
        "Authentication or authorisation",
        "Real-time collaboration",
        "Mobile application",
        "Neo4j integration (available but disabled by default)",
    ]
    for item in out_scope:
        bullet(doc, item)
    doc.add_paragraph()


def add_s6_functional(doc):
    section_heading(doc, "6", "Functional Requirements")

    def fr_table(sub, title, rows):
        sub_heading(doc, f"{sub}  {title}")
        add_table(doc, ["ID", "Requirement"], rows, col_widths=[0.7, 5.8])

    fr_table("6.1", "Document Upload and Graph Extraction", [
        ("FR-01", "System shall accept PDF, DOCX, and TXT documents"),
        ("FR-02", "System shall extract 15–20 nodes and 15–20 edges per document using Claude Sonnet"),
        ("FR-03", "Nodes shall be classified as: Process, Role, System, Policy, DataEntity, Event"),
        ("FR-04", "Extracted graph shall be rendered as an interactive force-directed network"),
        ("FR-05", "Each upload shall produce an independent graph with a unique graph_id (UUID)"),
        ("FR-06", "Document text stored in memory shall be capped at 15,000 characters"),
    ])
    fr_table("6.2", "Process Mining (Workflows tab)", [
        ("FR-07", "System shall provide a Celonis-style Process Mining surface as the primary content of the Workflows tab; legacy AS-IS / TO-BE narrative cards have been retired"),
        ("FR-08", "System shall expose GET /api/process-mining/{graph_id} returning KPIs, activities, transition edges, variants, conformance, and AI recommendations — pure SQL aggregation over process_steps + loan_applications + step_executions; no LLM calls"),
        ("FR-09", "Process Mining shall surface a KPI strip with Cases / Median TAT / Conformance % / SLA Breaches % / Role Mismatches % / Avg Loan, each tile carrying a drilldown subtitle (e.g. '8 disbursed · 3 active · 2 declined')"),
        ("FR-10", "Process Map shall render activities as absolutely-positioned SVG nodes; transition edges shall encode case count via stroke-width, with red for SLA-breach edges, amber dashed for rework loops"),
        ("FR-11", "System shall identify a single bottleneck activity by composite score (breach_count, role_mismatch_count, dwell/SLA ratio) and surface it via a BOTTLENECK pill plus the Bottleneck KPI tile on the Dashboard"),
        ("FR-12", "Variants tab shall list unique step sequences with frequency %, conformant/total counts, median TAT, swimlane detail and comparison vs Variant #1; clicking a variant updates the detail panel"),
        ("FR-13", "Execution Conformance tab shall show a fitness score, deviation patterns ranked by severity × case count, and a top-deviating-cases table (case ID / applicant / deviation / TAT / severity)"),
        ("FR-14", "Filter Cases popover shall accept categorical filters (loan type, status); applying re-fetches the endpoint with ?loan_types=…&statuses=… query params and recomputes all downstream KPIs / variants / conformance coherently"),
        ("FR-15", "Negative transition times computed from overlapping seed dates shall be clamped to ≥0; skipped-step detection shall use set-membership, not sequence position; out-of-sequence execution shall be its own deviation pattern"),
        ("FR-16", "Process Mining endpoint shall be uncached (recomputes on each fetch since it is fast SQL); the Optimise endpoint shall cache its Haiku result per graph_id"),
    ])
    fr_table("6.2.1", "AI Process Optimisation (Haiku)", [
        ("FR-16a", "System shall expose POST /api/process-mining/{graph_id}/optimise that runs a Haiku call against the current PM snapshot and returns 3–5 specific re-routing / SLA / role-assignment suggestions"),
        ("FR-16b", "Each suggestion shall include: title, rationale, expected_impact (qualitative or quantitative), effort (low/medium/high), and target_step (or empty for cross-cutting)"),
        ("FR-16c", "Optimise modal shall surface the suggestions with effort badges (low=green, medium=amber, high=red) and target-step references; the result shall be cached per graph_id so re-clicks are instant"),
    ])
    fr_table("6.3", "Gap Analysis", [
        ("FR-17", "System shall identify gaps: missing roles, policies, and systems in the process graph"),
        ("FR-18", "Each gap item shall have a severity: Critical / High / Medium / Low"),
        ("FR-19", "View in Graph shall highlight affected nodes on the knowledge graph"),
        ("FR-20", "System shall generate a remediation blueprint using Claude Haiku"),
        ("FR-21", "Gap analysis result shall be exportable to a plain-text file for offline review"),
    ])
    fr_table("6.4", "Pulse Recommendations", [
        ("FR-22", "System shall calculate rule-based health metrics from the knowledge graph"),
        ("FR-23", "System shall generate AI Strategic Insights (up to 5) using Claude Haiku, accessible via a button at the bottom of the Pulse drawer"),
        ("FR-24", "Pulse drawer shall slide in from a bell icon at the bottom of the sidebar"),
        ("FR-24a", "Pulse drawer shall surface a filter tab strip at the top (All + Critical / Warning / Info — only severities present in the items, with per-tab item counts)"),
        ("FR-24b", "Each Pulse card shall display a severity pill (CRITICAL / WARNING / INFO), a 'just now' timestamp, the title, a derived Source line (sourced from item.category — e.g. 'Pulse · Urgent'), and an action row with 'View in Graph' (filled brand pill) + 'Dismiss' (ghost pill)"),
        ("FR-24c", "Dismiss shall be local-state only — items remain hidden until the next /api/pulse/calculate fetch; no backend persistence"),
    ])
    fr_table("6.5", "Audit Check (was Conformance Checking — renamed)", [
        ("FR-25", "Standalone Conformance view shall be renamed Audit Check throughout the UI (sidebar nav, page title, score header, button labels, header chip) to disambiguate from Process Mining → Execution Conformance"),
        ("FR-26", "User shall upload an audit document (audit report, review log) — Audit Check is document-only; the live-operational-data ingestion path has been retired since that question is answered by Execution Conformance under Process Mining"),
        ("FR-27", "System shall score each eligible node as: Confirmed / Deviated / Not Found via a Sonnet call comparing the audit doc against the SOP graph"),
        ("FR-28", "System shall display an Audit Match Rate percentage with evidence quotes per finding"),
        ("FR-29", "Overlay modes shall allow filtering the always-visible knowledge graph to confirmed, deviated, or all nodes"),
        ("FR-30", "Event and Objective nodes shall be excluded from audit assessment"),
        ("FR-30a", "Audit Check shall NOT be cached on disk; same SOP+audit pair re-runs on every analyse call (so the AI moment is visibly real, not pre-cooked)"),
    ])
    fr_table("6.6", "Object Model Generation", [
        ("FR-31", "System shall generate a domain object model using BRD Authoring Standard §4"),
        ("FR-32", "Output shall include: Pydantic v2 Python code, JSON Schema, and an Entity Diagram (ERD)"),
        ("FR-33", "Object Model tab shall default to the Entity Diagram view (was Pydantic); tab order shall be Entity Diagram → Pydantic → JSON Schema"),
        ("FR-33a", "Topbar shall expose three action buttons — Pydantic and JSON Schema (jump-to-tab) and Regenerate (primary, brand-teal) — plus a meta line 'N ENTITIES · BRD §4'"),
        ("FR-33b", "Regenerate shall call POST /api/generate-object-model?force=true which busts both the in-memory store and the disk cache and re-runs the Sonnet call"),
        ("FR-34", "ERD shall lay out entity cards by foreign-key depth (column = number of inbound FK refs); cards shall be 240px-wide, absolutely positioned, with a flat brand-teal header containing a cube icon and a lowercase mono entity name"),
        ("FR-34a", "Each ERD field row shall show a PK / FK / AUDIT pill (or none) followed by the field name and type in mono; arrows between cards shall be solid brand-teal SVG paths"),
        ("FR-35", "Object Model shall be cached per graph_id (in-memory) and per doc_hash (on disk via demo_cache.json) — re-navigation is instant, but a customer-visible Regenerate is provided"),
        ("FR-35a", "If no graph has been extracted, system shall show an empty-state prompt"),
    ])
    fr_table("6.7", "Natural Language Querying", [
        ("FR-36", "User shall be able to type a plain-English question about the graph"),
        ("FR-37", "System shall return a direct answer plus highlighted relevant nodes"),
        ("FR-38", "Query history shall be maintained per graph_id for the session"),
        ("FR-39", "When the SQLite demo DB is seeded, the NLQ prompt shall include a compact OPERATIONAL DATA section so answers about reality (breach rates, role mismatches, cycle times) are grounded in actual data"),
    ])
    fr_table("6.8", "Dashboard", [
        ("FR-40", "Dashboard shall be a full-screen view (graph canvas hidden) that aggregates state from in-memory stores plus the Process Mining snapshot — no new Claude calls"),
        ("FR-41", "Dashboard shall display an Overall Health Score (0–100) as a weighted average of Coverage (40%), SLA Compliance (35% — derived from PM as 100 − breach_rate_pct), and Audit Conformance (25%); weights redistribute proportionally when sources are missing"),
        ("FR-42", "Dashboard hero shall show four click-through KPI tiles: Coverage / SLA / Conformance / Bottleneck (the Bottleneck tile replaces the legacy Total ROI tile and shows the worst PM activity with breach/role-mismatch detail)"),
        ("FR-43", "Dashboard shall show a graph composition bar chart by node type"),
        ("FR-44", "Dashboard shall show up to 5 Top Issues drawn from gap analysis (critical/warning) and PM deviation patterns — legacy workflow-SLA-step issues are gone"),
        ("FR-45", "Dashboard shall show up to 4 Top Hotspots — PM deviation patterns ranked by severity × case count (replaces the legacy 'Quick wins · ROI' card, which depended on the retired workflow narrative)"),
        ("FR-46", "Dashboard shall show up to 3 Top Deviating Cases — case ID, applicant name, deviation, severity badge — with click-through to Process Mining (replaces 'Automation Highlights')"),
        ("FR-47", "Dashboard shall show a Setup-Progress checklist: Graph extracted / Process Mining / Gap Analysis / Audit Check (was 'Workflows' / 'Conformance' — labels updated)"),
        ("FR-48", "When the SQLite demo DB is seeded, Dashboard shall include a Live Operational Data section with KPIs (applications, total value, avg cycle, step executions), a status breakdown, and top breached steps"),
        ("FR-49", "Dashboard shall present a hero CTA empty state when no graph is loaded"),
        ("FR-50", "Dashboard shall auto re-render after Process Mining data finishes loading post-upload, so hotspots and deviating-case sections populate without requiring a navigation"),
    ])
    fr_table("6.9", "Multi-Document Upload and Cross-Document Insights", [
        ("FR-51", "System shall accept multiple documents in a single upload request"),
        ("FR-52", "System shall produce ONE unified knowledge graph from a single Sonnet call against combined document text (no per-doc-then-merge)"),
        ("FR-53", "Combined input shall be capped at 15,000 characters; documents shall be truncated proportionally when the cap is exceeded"),
        ("FR-54", "Every node shall carry a 'sources' field listing the filename(s) it appeared in (multi-source = higher confidence)"),
        ("FR-55", "Upload UI shall display a combined-character counter (green / amber / red) that warns before the limit is exceeded"),
        ("FR-56", "Above the graph, the system shall display a colored chip per source document; clicking a chip shall filter the graph to nodes from that document"),
        ("FR-57", "Node detail panel shall display a Sources field listing the document filename(s) when present"),
        ("FR-58", "When two or more documents are uploaded, the system shall produce up to 3 Cross-Document Insights via a single Haiku call (categories: gap / inconsistency / missing / contradiction); cached per graph_id"),
        ("FR-59", "Single-document upload shall behave exactly as before — the chip row collapses to one chip and the Cross-Document Insights panel is hidden"),
    ])
    fr_table("6.10", "Live Operational Data (SQLite)", [
        ("FR-60", "System shall ship with a local SQLite demo DB seeded on server startup with realistic banking-loan operations data (process_steps, loan_applications, step_executions)"),
        ("FR-61", "System shall expose POST /api/data/summary returning aggregate stats (totals, status breakdown, average cycle time, top breached steps)"),
        ("FR-62", "System shall expose GET /api/data/step/{step_name} returning expected-vs-actual stats for a single step (execution count, breach rate, avg duration, role/system mismatches)"),
        ("FR-63", "Seeding shall be idempotent — re-running on a populated DB shall be a no-op"),
        ("FR-64", "No external service or credentials shall be required for the demo data store to function"),
    ])
    fr_table("6.11", "Executive Report", [
        ("FR-65", "System shall provide a header button 'Generate Report' that opens a self-contained, print-styled HTML page summarising every feature run for the current graph"),
        ("FR-66", "Report shall be composed entirely from existing in-memory stores and the SQLite DB; no new LLM calls shall be made to compose it (the AI Optimisation section, when present, is sourced from the cached optimise result)"),
        ("FR-67", "Report sections shall render conditionally — features that have not been run shall be omitted"),
        ("FR-68", "Report shall include the following sections in order: Cover (hero), Source Documents, Executive Summary, Top Issues, Cross-Document Insights (when applicable), Process Mining (KPIs + bottleneck + deviation patterns + top deviating cases — replaces legacy AS-IS/TO-BE Workflow Automation Opportunities), Gap Analysis + Blueprint, Audit Check (renamed from Conformance Results), Object Model (entity table), AI Optimisation Suggestions (when run), Footer"),
        ("FR-68a", "Report shall NOT include a standalone Live Operational Data section — that data is subsumed by Process Mining's KPI strip and bottleneck detail"),
        ("FR-69", "Report HTML shall include a print stylesheet (page-break rules, color preservation, hidden action bar) so the user can save to PDF via the browser's native Print dialog"),
        ("FR-69a", "Report typography shall use Inter for body text and IBM Plex Mono for numerics with letter-spacing -0.02em; section labels shall be 10px / 0.10em / brand-teal mono — matching the rest of the application's design tokens"),
        ("FR-69b", "Report cover (hero) shall use the wordmark 'Metafore · Discovery' (not legacy 'Metafore Works · Discovery Engine'); cover shall be on its own page in print via page-break-after"),
        ("FR-69c", "Print rules shall include @page { margin: 16mm 14mm } and page-break-before: always on each major section so each surface starts on its own page; even-row table backgrounds shall turn transparent in print"),
    ])
    fr_table("6.12", "Cache Behaviour & Demo Authenticity", [
        ("FR-70", "System shall maintain a disk cache (data/demo_cache.json) keyed by SHA-256 doc_hash for selected features: workflows, object_model, blueprint, pulse_ai, cross_doc"),
        ("FR-71", "demo_cache shall expose put / get / invalidate / stats functions; the invalidate helper shall be used by Object Model Regenerate (FR-33b)"),
        ("FR-72", "Graph extraction shall NOT be cached — each /api/upload re-runs the Sonnet call so the headline AI moment is visibly real on every demo replay"),
        ("FR-73", "Audit Check shall NOT be cached — each /conformance/analyse re-runs Sonnet so re-uploading the same audit document produces a visible AI moment"),
        ("FR-74", "Process Mining and Gap Analysis are pure SQL / rule-based — no LLM, no cache; recomputed on each fetch"),
        ("FR-75", "Object Model and Optimise shall remain cached (downstream of the headline extraction; navigation, not fresh-upload events)"),
    ])


def add_s7_nfr(doc):
    section_heading(doc, "7", "Non-Functional Requirements")
    add_table(doc,
        ["ID", "Category", "Requirement"],
        [
            ("NFR-01", "Performance",    "Graph extraction shall complete within 60 seconds for documents up to 20,000 characters"),
            ("NFR-02", "Performance",    "Natural language queries shall respond within 10 seconds"),
            ("NFR-03", "Usability",      "All features shall be accessible from a single-page browser UI with no page reloads"),
            ("NFR-04", "Reliability",    "Server shall auto-restart on code changes during development (uvicorn --reload)"),
            ("NFR-05", "Security",       "API key shall be stored in a local .env file and never committed to version control"),
            ("NFR-06", "Compatibility",  "Frontend shall work in Microsoft Edge without CDN dependency"),
            ("NFR-07", "Portability",    "Application shall run on any Windows machine with Python 3.11+ via double-click START.bat"),
            ("NFR-08", "Data isolation", "All session data is in-memory; restart clears all state"),
        ],
        col_widths=[0.7, 1.3, 4.5]
    )


def add_s8_object_model(doc):
    section_heading(doc, "8", "Object Model Standard — BRD Authoring Standard §4")
    body_text(doc,
        "All domain object models generated by the Discovery Engine must conform to the following rules. "
        "These rules apply to both the Pydantic code output and the JSON Schema output."
    )

    sub_heading(doc, "8.1  Field Naming")
    add_table(doc,
        ["Rule", "Specification"],
        [
            ("Case",    "snake_case only — no spaces, no camelCase"),
            ("Length",  "Maximum 63 characters per field name"),
            ("FK naming", "<parent_table>_id format — e.g. customer_id, account_id"),
        ],
        col_widths=[1.5, 5.0]
    )

    sub_heading(doc, "8.2  Mandatory Fields — Every Entity")
    body_text(doc, "Every entity must include these fields in this order:")
    add_table(doc,
        ["Position", "Field", "Type", "Constraint"],
        [
            ("1st", "id",         "UUID",     "Primary Key"),
            ("2nd", "created_at", "datetime", "Not Null — audit trail"),
        ],
        col_widths=[0.8, 1.5, 1.5, 2.7]
    )

    sub_heading(doc, "8.3  Audit Fields — Every Entity")
    body_text(doc, "Every entity must include these audit-only fields. They carry no FK constraint and must not appear in ERD relationship lines.")
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ("created_by",       "UUID", "AUDIT ONLY – no relationship"),
            ("last_changed_by",  "UUID", "AUDIT ONLY – no relationship"),
        ],
        col_widths=[2.0, 1.5, 3.0]
    )

    sub_heading(doc, "8.4  Enum Values")
    body_text(doc, "All enum values must be UPPERCASE: OPEN, IN_PROGRESS, APPROVED, REJECTED, CLOSED. Never lowercase.")

    sub_heading(doc, "8.5  Relationship Declaration Format")
    body_text(doc, "Relationships must use this exact string format:")
    add_note_box(doc, "entity_a (parent) → entity_b (child) via entity_b.fk_column = entity_a.id\n\n"
                      "Example: Customer (parent) → LoanApplication (child) via loan_application.customer_id = customer.id")

    sub_heading(doc, "8.6  JSON Schema Requirements")
    body_text(doc, 'The JSON Schema output must include an "entities" array (each with name and fields) '
                   'and a "relationships" array using the arrow format above.')


def add_s9_architecture(doc):
    section_heading(doc, "9", "System Architecture")

    sub_heading(doc, "9.1  Tech Stack")
    body_text(doc,
        "Backend: Python 3.11+, FastAPI, Uvicorn, anthropic SDK, python-docx (DOCX parsing — extracts paragraphs and table cells), "
        "PyMuPDF (PDF parsing), python-dotenv, sqlite3 (built-in, used for the local demo DB; no external service required). "
        "Frontend: Vanilla JavaScript + HTML + CSS, vis-network (self-hosted), no build step. "
        "The Discovery Engine ships as a single-file FastAPI app plus static frontend assets — no database, message queue, or auth provider required to run."
    )
    body_text(doc, "Architecture details:")
    add_table(doc,
        ["Layer", "Technology", "Notes"],
        [
            ("AI",                "Anthropic Claude Sonnet 4.6 + Haiku 4.5",  "Tool-use API — structured output guaranteed"),
            ("Backend",           "Python 3.11 · FastAPI 0.115 · Uvicorn 0.32", "Port 8083, launched via --app-dir"),
            ("Document parsing",  "PyMuPDF (PDF) · python-docx (DOCX)",         "Built-in for TXT"),
            ("Graph DB",          "Neo4j 5",                                     "Optional — disabled by default"),
            ("Frontend",          "Vanilla JS · vis-network 9.1.9",              "Self-hosted; no build step required"),
            ("Fonts",             "GT America · IBM Plex Mono · Inter",           "Self-hosted OTF/TTF"),
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    sub_heading(doc, "9.2  API Endpoints")
    add_table(doc,
        ["Endpoint", "Method", "Feature"],
        [
            ("/api/upload",                                "POST", "Multi-document upload + unified graph extraction"),
            ("/api/graph/{graph_id}/sources",              "GET",  "List source documents for a graph"),
            ("/api/graph/{graph_id}/cross-doc-insights",   "POST", "Cross-document gaps and inconsistencies (Haiku, cached)"),
            ("/api/workflows/generate",                    "POST", "Workflow bundle (workflow + ROI + automation + variants)"),
            ("/api/workflows/{graph_id}",                  "GET",  "Retrieve cached workflows"),
            ("/api/generate-object-model",                 "POST", "Object model generation"),
            ("/api/query/natural-language",                "POST", "Natural language query (with SQLite operational context)"),
            ("/api/query/history/{graph_id}",              "GET",  "NLQ history"),
            ("/api/gap-analysis/calculate",                "POST", "Gap analysis"),
            ("/api/gap-analysis/blueprint",                "POST", "Remediation blueprint"),
            ("/api/pulse/calculate",                       "POST", "Pulse health metrics"),
            ("/api/pulse/ai-recommendations",              "POST", "AI pulse recommendations"),
            ("/api/data/summary",                          "POST", "SQLite operational data — aggregate stats"),
            ("/api/data/step/{step_name}",                 "GET",  "SQLite operational data — per-step expected-vs-actual"),
            ("/api/report/{graph_id}",                     "GET",  "Executive HTML report (browser print to PDF)"),
            ("/api/health",                                "GET",  "Health probe (neo4j_enabled, demo_db_seeded)"),
            ("/conformance/upload",                        "POST", "Evidence document upload (file or synthetic from live data)"),
            ("/conformance/analyse",                       "POST", "Conformance analysis"),
            ("/conformance/{graph_id}/latest",             "GET",  "Latest conformance result"),
        ],
        col_widths=[3.4, 0.7, 2.4]
    )

    sub_heading(doc, "9.3  AI Model Routing")
    add_table(doc,
        ["Feature", "Model", "Max Tokens", "Reason"],
        [
            ("Graph extraction (single or multi-doc)", "Claude Sonnet 4.6", "8,096",  "Highest accuracy on entity/relationship extraction"),
            ("Workflow bundle (workflow + ROI + automation + variants)", "Claude Sonnet 4.6", "32,000", "One consolidated call replaces former 7 calls; system prompt + tool schema cached (ephemeral)"),
            ("Object model",                                "Claude Sonnet 4.6", "4,096",  "Code generation quality"),
            ("Conformance analysis",                        "Claude Sonnet 4.6", "6,000",  "Precise evidence matching"),
            ("Natural language query",                      "Claude Haiku 4.5",  "800",    "Simple Q&A; operational data injected into prompt"),
            ("Blueprint",                                   "Claude Haiku 4.5",  "800",    "Short summary"),
            ("Pulse AI recommendations",                    "Claude Haiku 4.5",  "2,048",  "Structured list"),
            ("Cross-document insights",                     "Claude Haiku 4.5",  "1,500",  "Up to 3 cross-doc gaps, cached per graph_id"),
        ],
        col_widths=[2.6, 1.8, 0.8, 1.3]
    )


def add_s10_data_model(doc):
    section_heading(doc, "10", "Data Model")

    sub_heading(doc, "10.1  In-Memory Stores")
    body_text(doc, "Most stores are keyed by graph_id (UUID) and reset on server restart. _workflow_cache is keyed by document hash to enable cross-session caching. There is no persistent storage for graph state — only the SQLite demo DB persists between sessions.")
    add_table(doc,
        ["Store", "Key", "Value"],
        [
            ("_graph_store",          "graph_id",    "Nodes, edges, graph_id (each node carries a 'sources' field)"),
            ("_doc_store",            "graph_id",    "Combined document text (max 15,000 chars)"),
            ("_doc_sources_store",    "graph_id",    "List of {filename, word_count} for source documents"),
            ("_hash_store",           "graph_id",    "SHA-256 hash of uploaded raw bytes (workflow cross-session cache key)"),
            ("_workflow_store",       "graph_id",    "Workflow bundle (workflows + ROI + automation + variants)"),
            ("_workflow_cache",       "doc_hash",    "Cross-session workflow cache (same content → instant reuse)"),
            ("_cross_doc_store",      "graph_id",    "Cross-document insights result"),
            ("_object_model_store",   "graph_id",    "Object model result"),
            ("_query_history",        "graph_id",    "NLQ history entries"),
            ("_gap_store",            "graph_id",    "Gap analysis result"),
            ("_blueprint_store",      "graph_id",    "Blueprint result"),
            ("_pulse_store",          "graph_id",    "Pulse items"),
            ("_pulse_ai_store",       "graph_id",    "AI recommendations"),
            ("_evidence_store",       "evidence_id", "Evidence document"),
            ("_conformance_store",    "evidence_id", "Conformance result"),
            ("_conformance_latest",   "graph_id",    "Latest evidence_id"),
        ],
        col_widths=[2.2, 1.3, 3.0]
    )

    sub_heading(doc, "10.1a  SQLite Demo DB (persistent local file)")
    body_text(doc, "Auto-seeded on server startup at backend/data/demo.db. Holds banking-loan operational data the dashboard, workflows, conformance, and NLQ overlay onto the SOP-extracted graph.")
    add_table(doc,
        ["Table", "Purpose"],
        [
            ("process_steps",   "SOP-defined steps with expected_role, expected_system, sla_hours"),
            ("loan_applications","13 sample applications with status (disbursed / underwriting / pending / declined) and amount"),
            ("step_executions", "~71 actual step executions with started_at, completed_at, actual_role, actual_system, status (completed / breached / skipped / in_progress)"),
        ],
        col_widths=[2.0, 4.5]
    )

    sub_heading(doc, "10.2  Node Types")
    add_table(doc,
        ["Type", "Colour", "Represents"],
        [
            ("Process",    "Teal",   "Business processes and activities"),
            ("Role",       "Orange", "People, teams, job functions"),
            ("System",     "Blue",   "Applications, platforms, databases"),
            ("Policy",     "Red",    "Rules, regulations, compliance requirements"),
            ("DataEntity", "Purple", "Data objects, records, documents"),
            ("Event",      "Yellow", "Triggers, milestones, notifications"),
        ],
        col_widths=[1.3, 1.2, 4.0]
    )


def add_s11_ui(doc):
    section_heading(doc, "11", "User Interface Requirements")
    add_table(doc,
        ["ID", "Requirement"],
        [
            ("UI-01", "Navigation via a 200px labeled sidebar with five items: Graph, Workflows, Gap Analysis, Conformance, Object Model"),
            ("UI-02", "Left panel width varies per view: Graph 35%, Workflows 60%, Gap 50%, Conformance 50%, Object Model 55%"),
            ("UI-03", "Right graph canvas shall always be visible regardless of the active left panel"),
            ("UI-04", "Sidebar top shows Metafore Works logo and text; header shows DISCOVERY ENGINE only"),
            ("UI-05", "Colour palette follows the Metafore Works design system (primary #036868)"),
            ("UI-06", "All node types in the graph shall use circle shapes only — no mixed shapes"),
            ("UI-07", "Object model panel shall show a loading spinner during API calls and empty-state when no graph is loaded"),
        ],
        col_widths=[0.7, 5.8]
    )


def add_s12_samples(doc):
    section_heading(doc, "12", "Sample Demo Documents")
    add_table(doc,
        ["Document", "Purpose", "Expected Output"],
        [
            ("commercial_banking_loan_sop.docx", "Source SOP for extraction",       "~25 nodes: 8 processes, 6 roles, 4 systems, 3+ policies"),
            ("loan_audit_report.docx",           "Evidence for conformance check",  "~47% conformance: 9 confirmed, 10 deviated, 2 not found"),
        ],
        col_widths=[2.2, 1.8, 2.5]
    )
    body_text(doc,
        "Key conformance deviations surfaced by the audit report: AML check sequencing (Critical), "
        "property appraisal before completeness check, sanctioning timeline breach, "
        "DTI calculation assigned to incorrect role."
    )


def add_s13_constraints(doc):
    section_heading(doc, "13", "Constraints and Assumptions")
    add_table(doc,
        ["Type", "Detail"],
        [
            ("Constraint", "Requires a valid Anthropic API key with sufficient credits"),
            ("Constraint", "Windows only — batch scripts use .bat syntax"),
            ("Constraint", "Python 3.11+ must be installed with PATH configured"),
            ("Constraint", "All data is session-scoped; no persistence between server restarts"),
            ("Assumption", "Single user at a time — no concurrent session isolation"),
            ("Assumption", "Documents are in English"),
            ("Assumption", "Network access is available for initial pip install"),
        ],
        col_widths=[1.3, 5.2]
    )


def add_s14_glossary(doc):
    section_heading(doc, "14", "Glossary")
    add_table(doc,
        ["Term", "Definition"],
        [
            ("BRD",          "Business Requirements Document"),
            ("ERD",          "Entity Relationship Diagram"),
            ("SOP",          "Standard Operating Procedure"),
            ("NLQ",          "Natural Language Query"),
            ("graph_id",     "UUID assigned to each extracted knowledge graph"),
            ("evidence_id",  "UUID assigned to each uploaded conformance evidence document"),
            ("Pulse",        "Health score and recommendations derived from the knowledge graph"),
            ("Blueprint",    "AI-generated remediation plan for identified gaps"),
            ("Conformance",  "Assessment of how closely an evidence document matches the source process graph"),
            ("AS-IS",        "Current state of a business process as documented"),
            ("TO-BE",        "Target future state of a process with recommended improvements"),
        ],
        col_widths=[1.5, 5.0]
    )


# ══════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    add_cover(doc)
    add_meta_table(doc)
    add_s1_summary(doc)
    add_s2_context(doc)
    add_s3_objectives(doc)
    add_s4_stakeholders(doc)
    add_s5_scope(doc)
    add_s6_functional(doc)
    add_s7_nfr(doc)
    add_s8_object_model(doc)
    add_s9_architecture(doc)
    add_s10_data_model(doc)
    add_s11_ui(doc)
    add_s12_samples(doc)
    add_s13_constraints(doc)
    add_s14_glossary(doc)

    path = OUT / "BRD_Discovery_Engine.docx"
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build()
