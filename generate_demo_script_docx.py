"""Generate a professional DOCX version of the Demo Script."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent

BRAND      = RGBColor(0x03, 0x68, 0x68)
BRAND_DARK = RGBColor(0x02, 0x4F, 0x4F)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TEAL_LIGHT = RGBColor(0xE6, 0xF4, 0xF4)
TEAL_MID   = RGBColor(0xA7, 0xD4, 0xD4)
TEXT_DARK  = RGBColor(0x1F, 0x2D, 0x2D)
TEXT_MID   = RGBColor(0x4A, 0x68, 0x68)
QUOTE_BG   = RGBColor(0xF0, 0xFA, 0xFA)
STEP_BG    = RGBColor(0xF7, 0xFC, 0xFC)
AMBER      = RGBColor(0xD9, 0x77, 0x06)
GREEN      = RGBColor(0x16, 0xA3, 0x4A)
GREY       = RGBColor(0x6B, 0x72, 0x80)


def set_cell_bg(cell, rgb):
    hex_color = str(rgb) if not isinstance(rgb, str) else rgb
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def cell_para(cell, text, bold=False, size=9, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.paragraphs[0].alignment = align
    return cell.paragraphs[0]


def add_run(para, text, bold=False, italic=False, size=None, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


# ── Cover / header banner ─────────────────────────────────────────────

def add_cover(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, BRAND_DARK)

    p1 = cell.paragraphs[0]
    p1.clear()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("METAFORE")
    r.font.color.rgb = TEAL_MID
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.letter_spacing = Pt(2)  # may not render in all viewers

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Discovery Engine")
    r2.font.color.rgb = WHITE
    r2.font.size = Pt(22)
    r2.font.bold = True

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Demo Script")
    r3.font.color.rgb = TEAL_MID
    r3.font.size = Pt(13)
    r3.font.bold = False

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Approx. 12 – 15 minutes   |   http://localhost:8083")
    r4.font.color.rgb = TEAL_MID
    r4.font.size = Pt(8.5)

    doc.add_paragraph()  # spacer


# ── Quick-info box at the top ─────────────────────────────────────────

def add_info_box(doc):
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    rows_data = [
        ("URL",               "http://localhost:8083"),
        ("SOP Document",      "sample_docs/commercial_banking_loan_sop.docx"),
        ("Audit Document",    "sample_docs/loan_audit_report.docx"),
    ]
    for i, (lbl, val) in enumerate(rows_data):
        set_cell_bg(table.rows[i].cells[0], TEAL_LIGHT)
        cell_para(table.rows[i].cells[0], lbl, bold=True, size=9, color=BRAND_DARK)
        cell_para(table.rows[i].cells[1], val, size=9)
    doc.add_paragraph()


# ── Section heading ───────────────────────────────────────────────────

def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = BRAND
    # top border via paragraph XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), str(BRAND))
    pBdr.append(top)
    pPr.append(pBdr)
    return p


# ── Opening quote box ─────────────────────────────────────────────────

def add_quote_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, QUOTE_BG)
    # left border accent via tcBorders
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), str(BRAND))
    tcBorders.append(left)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_DARK
    doc.add_paragraph()


# ── Step card ─────────────────────────────────────────────────────────

def add_step_card(doc, step_num, title, duration, actions, narrations,
                  callouts=None, tips=None):
    """
    actions   : list of str (numbered action items)
    narrations: list of (after_action_idx, narration_text)  — 0-based index
    callouts  : optional list of (bullet_label, bullet_text) highlighted points
    tips      : optional list of str tip lines
    """
    # ── Header row ──
    header_table = doc.add_table(rows=1, cols=2)
    lc = header_table.cell(0, 0)
    rc = header_table.cell(0, 1)
    set_cell_bg(lc, BRAND)
    set_cell_bg(rc, BRAND)
    lc.width = Inches(4.5)
    rc.width = Inches(2.0)

    p_l = lc.paragraphs[0]
    p_l.clear()
    add_run(p_l, f"Step {step_num}  ", bold=True, size=10, color=WHITE)
    add_run(p_l, title, bold=False, size=10, color=TEAL_MID)

    p_r = rc.paragraphs[0]
    p_r.clear()
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_r, duration, bold=False, size=9, color=TEAL_MID)

    # ── Body ──
    narr_map = {idx: text for idx, text in narrations}

    for i, action in enumerate(actions):
        # action row
        body_table = doc.add_table(rows=1, cols=2)
        body_table.style = 'Table Grid'
        nc = body_table.cell(0, 0)
        ac = body_table.cell(0, 1)
        set_cell_bg(nc, TEAL_LIGHT)
        set_cell_bg(ac, STEP_BG)
        nc.width = Inches(0.35)
        ac.width = Inches(6.15)
        cell_para(nc, str(i + 1), bold=True, size=9, color=BRAND,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(ac, action, size=9)

        # narration after this action?
        if i in narr_map:
            nq_table = doc.add_table(rows=1, cols=1)
            nc2 = nq_table.cell(0, 0)
            set_cell_bg(nc2, QUOTE_BG)
            _add_left_border(nc2, str(BRAND))
            p = nc2.paragraphs[0]
            p.clear()
            add_run(p, f'Say:  "{narr_map[i]}"', italic=True, size=9,
                    color=BRAND_DARK)

    # callouts
    if callouts:
        co_table = doc.add_table(rows=len(callouts), cols=2)
        co_table.style = 'Table Grid'
        for i, (lbl, txt) in enumerate(callouts):
            set_cell_bg(co_table.rows[i].cells[0], TEAL_LIGHT)
            set_cell_bg(co_table.rows[i].cells[1], STEP_BG)
            cell_para(co_table.rows[i].cells[0], lbl, bold=True, size=8.5,
                      color=BRAND_DARK)
            cell_para(co_table.rows[i].cells[1], txt, size=8.5)

    # tips
    if tips:
        tip_table = doc.add_table(rows=1, cols=1)
        tc2 = tip_table.cell(0, 0)
        set_cell_bg(tc2, RGBColor(0xFF, 0xFB, 0xEB))
        _add_left_border(tc2, str(AMBER))
        p = tc2.paragraphs[0]
        p.clear()
        add_run(p, "Tips:  ", bold=True, size=8.5, color=AMBER)
        add_run(p, "  |  ".join(tips), size=8.5, color=RGBColor(0x78, 0x35, 0x00))

    doc.add_paragraph()


def _add_left_border(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)


# ── Quick reference table ─────────────────────────────────────────────

def add_quick_ref(doc):
    section_heading(doc, "Quick Reference — What Each Feature Shows")
    doc.add_paragraph()

    rows = [
        ("Graph Extraction",     "Turns documents into structured knowledge",                         "Yes — Sonnet",       "~30–40s"),
        ("Gap Analysis",         "Structural audit — no AI needed",                                   "No",                 "Instant"),
        ("Blueprint",            "Actionable improvement plan",                                       "Yes — Haiku",        "~7s"),
        ("Process Mining",       "Celonis-class operational mining — bottleneck, fitness, deviations", "No (pure SQL)",      "Instant"),
        ("AI Optimise",          "Re-routing / SLA / role-assignment suggestions",                    "Yes — Haiku",        "~6s"),
        ("Audit Check",          "Audit document vs SOP — Claude reads prose evidence",               "Yes — Sonnet",       "~30–40s"),
        ("Object Model",         "Entity Diagram + Pydantic + JSON Schema (BRD §4)",                  "Yes — Sonnet",       "~30s"),
        ("Pulse Drawer",         "Severity-filtered action queue with Dismiss",                       "No / Haiku (insights)", "Instant"),
        ("Dashboard",            "Bird's-eye view sourced from PM + Gap + Audit",                     "No",                 "Instant"),
        ("Executive Report",     "Print-ready HTML — all sections, Metafore brand",                   "No",                 "Instant"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    headers = ("Feature", "Key Message", "AI Call?", "Wait Time")
    for c, h in enumerate(headers):
        set_cell_bg(table.rows[0].cells[c], BRAND)
        cell_para(table.rows[0].cells[c], h, bold=True, color=WHITE, size=9)

    for i, (feat, msg, ai, wait) in enumerate(rows):
        r = table.rows[i + 1]
        bg = TEAL_LIGHT if i % 2 == 0 else WHITE
        set_cell_bg(r.cells[0], bg)
        set_cell_bg(r.cells[1], bg)
        set_cell_bg(r.cells[2], bg)
        set_cell_bg(r.cells[3], bg)
        cell_para(r.cells[0], feat, bold=True, size=9, color=BRAND_DARK)
        cell_para(r.cells[1], msg,  size=9)
        ai_color = GREEN if ai.startswith("Yes") else GREY
        cell_para(r.cells[2], ai,   size=9, color=ai_color)
        cell_para(r.cells[3], wait, size=9)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    add_cover(doc)
    add_info_box(doc)

    # Opening
    section_heading(doc, "Opening Line  (30 seconds)")
    doc.add_paragraph()
    add_quote_box(doc,
        "Before you can build, automate, or evolve an enterprise application, you have to understand "
        "the organisation. Today, that takes consultants months and tools like Celonis cost half a "
        "million in integration before you see a single insight. The Discovery Engine reads your "
        "documents and gives you a living knowledge graph in minutes. Let me show you on a real "
        "banking SOP."
    )

    # ── Steps ──────────────────────────────────────────────────────────
    section_heading(doc, "Demo Steps")
    doc.add_paragraph()

    add_step_card(doc,
        step_num=1,
        title="Upload the Banking SOP & Knowledge Graph",
        duration="≈ 2 min",
        actions=[
            "Open http://localhost:8083 — you land on the Knowledge Graph tab",
            "Click Choose Files → select commercial_banking_loan_sop.docx",
            "Click Extract Knowledge Graph",
            "While loading (~30–40s) narrate the AI moment (see right →)",
            "When the graph renders, walk through the colour callouts (below)",
            "Click any node → inspector on the right shows the source-text quote",
        ],
        narrations=[
            (3, "This is a Loan Origination SOP — the kind of document that sits in a SharePoint and never gets read. "
                "The Discovery Engine is extracting every process, role, system, policy, and data entity from it automatically. "
                "This is the first AI moment — it's actually reading."),
            (5, "Every circle is an entity, every line is a relationship. The bank's lending process — extracted in under a minute."),
            (6, "Click any node, and the inspector shows you exactly where it came from in the original document. "
                "Nothing hallucinated, everything traceable."),
        ],
        callouts=[
            ("Teal circles",   "Processes — KYC Verification, Credit Check, Underwriting Review"),
            ("Orange circles", "Roles — Loan Officer, Credit Analyst, Credit Officer"),
            ("Blue circles",   "Systems — KYC Platform, Credit Bureau Gateway, Core Banking"),
            ("Red circles",    "Policies — AML, DTI threshold, sanctioning timeline"),
        ],
    )

    add_step_card(doc,
        step_num=2,
        title="Gap Analysis",
        duration="≈ 1 min",
        actions=[
            "Click the Gap Analysis nav (full-screen, graph hidden)",
            "Click Analyse Graph — results appear instantly (no AI call)",
            "Point to the coverage score and the breakdown",
            "Click a failing check → click View in Graph",
        ],
        narrations=[
            (1, "This is a structural audit of the graph. Does every process have an owner? Is every policy enforced? "
                "Are there orphans? It runs instantly — no Claude call, just rules."),
            (3, "The score tells us how complete the SOP itself is. Anything below 70 means the document has real gaps — "
                "processes without assigned roles, policies floating in space."),
            (4, "View in Graph jumps to the affected nodes. The graph becomes a heatmap of where the SOP is weak."),
        ],
    )

    add_step_card(doc,
        step_num=3,
        title="Generate Blueprint",
        duration="≈ 45 sec",
        actions=[
            "Still in Gap Analysis, click Generate Blueprint",
            "Wait ~7s (Haiku)",
            "Read out one of the next steps",
        ],
        narrations=[
            (1, "The Blueprint turns the gaps into an actionable improvement plan — what to fix, what to add, "
                "what documents to upload next. This is the output you hand to the process owner."),
        ],
    )

    add_step_card(doc,
        step_num=4,
        title="Process Mining (the headline)",
        duration="≈ 3 min",
        actions=[
            "Click Workflows nav — you land on Process Mining (full screen, graph hidden)",
            "FIRST: state the live-data caveat clearly (see narration 1)",
            "Walk through the KPI strip — Cases, Median TAT, Conformance, SLA Breaches, Role Mismatches, Avg Loan",
            "Process Map tab: point to KYC Verification (BOTTLENECK pill), edge thickness = case volume, red = SLA-breach edges, amber dashed = rework loops",
            "Click the bottleneck node → inspector shows wait-time chart, top causes, breach counts",
            "Variants tab: click a variant card → swimlane updates; show Sunita's case (skipped Underwriting Review entirely)",
            "Execution Conformance tab: point to fitness 0.46 — 6 of 13 cases conform; deviation patterns sorted by severity × case count",
            "Click Filter Cases button → check Commercial only → Apply (5 cases, fitness jumps to 60%); Clear to reset",
            "Click Optimise button → modal opens with Haiku-generated re-routing suggestions, effort badges, target steps",
        ],
        narrations=[
            (1, "Quick context — what you're about to see uses operational event-log data. "
                "For this demo we've pre-seeded a representative loan-origination dataset (13 cases) so you can see the mining surface live. "
                "In production this is wired from your event logs in days, not the six months Celonis takes."),
            (3, "13 cases mined, 7 activities, 8 days median TAT, 9% SLA breach rate, 46% conformance fitness. "
                "Every number is real, derived from operational data."),
            (4, "KYC Verification is flagged as the bottleneck — 2 SLA breaches across the runs."),
            (6, "Seven distinct execution paths through this process. Sunita Iyer's case actually skipped Underwriting Review entirely. "
                "We caught it from the data, not the SOP."),
            (7, "Fitness 0.46 — only 6 of 13 cases conform. The deviation patterns are listed by severity — "
                "skipped Underwriting (critical), wrong-role Credit Check (high), SLA-breached KYC (medium)."),
            (8, "You can slice by loan type. Commercial-only — 5 cases, fitness jumps to 60%. Retail loans are where most deviations live."),
            (9, "And here's the Haiku call. Given the bottleneck and the deviation patterns, the AI proposes specific re-routing rules and SLA targets — "
                "auto-route low-risk KYC, enforce mandatory Underwriting gates. Each suggestion has effort and impact estimates."),
        ],
        tips=[
            "Live-data caveat is mandatory — say it BEFORE you reveal numbers, not after",
            "Sunita / James / BlueRiver are real names in the seed data — let them land",
        ],
    )

    add_step_card(doc,
        step_num=5,
        title="Audit Check  (was Conformance — renamed)",
        duration="≈ 2 min",
        actions=[
            "Click Audit Check nav (was 'Conformance' — renamed for clarity)",
            "Click Choose File → select loan_audit_report.docx",
            "Click Run Audit Check →",
            "Wait ~30–40s (Sonnet — second AI moment, narrate while loading)",
            "Point to the audit match score (~30%); click a red deviated card → click View in Graph",
            "Toggle All / Deviated / Confirmed overlays",
        ],
        narrations=[
            (1, "Process Mining compared the SOP to your operational data. "
                "Audit Check compares the SOP to an audit document — a different question, answered by Claude reading prose."),
            (4, "This is a Q3 2024 audit report from the same bank. Claude is reading every node in the SOP graph and asking: "
                "did the audit confirm this happened? deviate? not mention it?"),
            (5, "30% audit match. Three nodes confirmed, six deviated, one no-evidence. The audit found that AML screening was performed "
                "AFTER credit pull — a critical sequencing breach the SOP forbids."),
            (6, "Every deviation links to the specific node and quotes the exact evidence from the audit. "
                "Different overlay views for different audiences — compliance team sees deviated, process owner sees confirmed."),
        ],
    )

    add_step_card(doc,
        step_num=6,
        title="Object Model",
        duration="≈ 1 min",
        actions=[
            "Click Object Model nav (full-screen, Entity Diagram tab default)",
            "Wait while it loads (Sonnet, ~30s)",
            "Point to the diagram — entities laid out by FK depth (customer → loan_application → credit_assessment / sanction)",
            "Click Pydantic / JSON Schema buttons in the topbar to jump tabs",
            "Click Regenerate → re-runs Sonnet (cache busted)",
        ],
        narrations=[
            (1, "The Object Model is the bridge from understanding to building. The same knowledge graph, "
                "expressed as a BRD-compliant entity model — Pydantic, JSON Schema, ERD. "
                "This is what the Application Language Model consumes upstream."),
            (3, "Layered by foreign-key depth. Brand-teal arrows are relationships. "
                "Click Pydantic or JSON Schema for the generated code."),
            (5, "Regenerate forces a fresh run if the graph has changed."),
        ],
    )

    add_step_card(doc,
        step_num=7,
        title="Dashboard — bird's-eye view",
        duration="≈ 1 min",
        actions=[
            "Click Dashboard nav (full-screen)",
            "Walk through: health score, four KPI tiles (Coverage / SLA / Conformance / Bottleneck), Top Hotspots, Top Deviating Cases",
            "Graph composition bar, Live Operational Data section, Setup-progress checklist",
            "Click the Bottleneck KPI tile or any Top Hotspot row → drills through to source tab",
        ],
        narrations=[
            (1, "This is the operational health of the discovery, in one view. Health score, four KPI tiles, "
                "top hotspots from Process Mining, top deviating cases by name, graph composition, live ops, completeness checklist."),
            (4, "Every tile and row drills through to the source tab. The dashboard isn't a report — it's the hub."),
        ],
    )

    add_step_card(doc,
        step_num=8,
        title="Pulse Drawer",
        duration="≈ 45 sec",
        actions=[
            "Click the bell icon at the bottom of the sidebar",
            "Drawer slides in — filter tabs at the top (All / Critical / Warning / Info with counts)",
            "Each card: severity pill, 'just now' timestamp, title, derived Source line, View in Graph + Dismiss buttons",
            "Click View in Graph on a critical item",
        ],
        narrations=[
            (1, "Pulse surfaces what needs attention now. Critical / Warning / Info filter at the top. "
                "Each card has the issue, source, and two actions — View in Graph or Dismiss."),
            (4, "Closes the drawer, jumps to the graph, highlights the affected nodes, banners the source. "
                "The agent's recommendation lands at the exact spot it applies."),
        ],
    )

    add_step_card(doc,
        step_num=9,
        title="Generate Executive Report  (the finale)",
        duration="≈ 1 min",
        actions=[
            "Click 'Generate Report' in the top-right of the header",
            "A new tab opens with the print-styled HTML report",
            "Walk through the sections quickly: cover, executive summary, top issues, Process Mining, Audit Check, Object Model, AI Optimisation",
            "Click Save as PDF / Print (or Ctrl+P) → show that each major section starts on its own page",
        ],
        narrations=[
            (1, "Cover page, Executive Summary with health score, Top Issues, Process Mining findings, Audit Check results, "
                "Object Model summary, AI Optimisation Suggestions. Self-contained HTML — they can save to PDF from the browser print dialog. "
                "No new AI calls — it's composed from what we already produced."),
            (3, "Each major section starts on its own page. The Metafore wordmark is on the cover. Ready for the boardroom."),
        ],
        tips=[
            "If you only have one slide to leave behind, this PDF is it",
            "Button is disabled until a graph has been extracted — extract first",
        ],
    )

    # Closing
    section_heading(doc, "Closing Line")
    doc.add_paragraph()
    add_quote_box(doc,
        "Twelve minutes. From a Word document to a knowledge graph, an operational process model with named deviations, "
        "an audit comparison, a generated data model, and a printable executive report. Today this is reading documents. "
        "Tomorrow — once you wire your event logs and systems of record — it's the same engine, the same graph, the same insights, "
        "but now living and self-evolving. That's the Metafore Discovery Engine. Slide 8 of our deck. The wedge into the EGI platform."
    )

    # Tips
    section_heading(doc, "Presenter Tips & Risks")
    doc.add_paragraph()
    tips_data = [
        ("Live-data question = #1 risk",
         "If asked 'can it read our SAP/Salesforce live?' — answer: 'Document-driven today; connector to live systems is days, not the months Celonis takes. We pre-seeded operational data here so you can see the mining surface.' Pre-empt the question; don't dodge."),
        ("If a Claude call fails or is slow",
         "Skip to the next feature — Gap Analysis, Process Mining, Pulse, and Dashboard all work without LLM calls."),
        ("Graph zoom",
         "The graph looks more impressive zoomed into dense clusters — use the scroll wheel."),
        ("Audit Check vs Execution Conformance",
         "Two different questions. Audit Check = doc-vs-doc (Claude reads prose). Execution Conformance (in PM tab) = data-vs-doc (pure SQL). Don't conflate them."),
        ("Named deviations land",
         "Sunita Iyer skipped Underwriting; James Patel was wrong-role on Approval; BlueRiver got a wrong-role Credit Check. Use the names."),
        ("Reset between demos",
         "Click Reset — clears all state cleanly in one click."),
        ("Cache behaviour",
         "Graph extraction and Audit Check are NOT cached (so the AI moments are visibly real every time). Object Model, Blueprint, Pulse AI, Cross-doc, and Optimise ARE cached."),
    ]
    table = doc.add_table(rows=len(tips_data), cols=2)
    table.style = 'Table Grid'
    for i, (lbl, tip) in enumerate(tips_data):
        bg = TEAL_LIGHT if i % 2 == 0 else WHITE
        set_cell_bg(table.rows[i].cells[0], bg)
        set_cell_bg(table.rows[i].cells[1], bg)
        cell_para(table.rows[i].cells[0], lbl, bold=True, size=9, color=BRAND_DARK)
        cell_para(table.rows[i].cells[1], tip, size=9)
    doc.add_paragraph()

    # Quick ref
    add_quick_ref(doc)

    path = OUT / "DEMO_SCRIPT.docx"
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build()
